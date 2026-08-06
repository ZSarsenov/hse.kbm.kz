"""Генератор коммерческого предложения для АО «Каражанбасмунай»."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BRAND_BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, color=BRAND_BLUE, space_before=6, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_paragraph(doc, text, size=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(size)


def make_doc():
    doc = Document()

    # Глобальный стиль
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Поля
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ===== ШАПКА =====
    header = doc.add_table(rows=1, cols=2)
    header.autofit = True
    left = header.cell(0, 0)
    right = header.cell(0, 1)

    p = left.paragraphs[0]
    r = p.add_run('[ТОО «___»]')
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BRAND_BLUE

    info_lines = [
        '[Адрес: г. ___, ул. ___, д. ___]',
        '[БИН: ___]',
        '[ИИК: ___ / БИК: ___]',
        '[Тел.: +7 (___) ___-__-__]',
        '[E-mail: ___@___.kz]',
    ]
    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, line in enumerate(info_lines):
        if i > 0:
            p2 = right.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p2.add_run(line)
        run.font.size = Pt(9)
        run.font.color.rgb = GREY

    doc.add_paragraph()

    # ===== ИСХ. НОМЕР И ДАТА =====
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = ''
    p_out = t.cell(0, 0).paragraphs[0]
    p_out.add_run('Исх. № [___]   от  «___» _________ 2025 г.').font.size = Pt(10)

    # ===== АДРЕСАТ =====
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    addressee_lines = [
        ('АО «Каражанбасмунай»', True),
        ('Генеральному директору', True),
        ('г-ну Сәрсенбай Нұрпейіс Мизанбайұлы', True),
        ('', False),
        ('Первому заместителю генерального директора', True),
        ('г-ну Ван Вэй', True),
    ]
    for i, (line, bold) in enumerate(addressee_lines):
        if i > 0:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(11)

    doc.add_paragraph()

    # ===== ЗАГОЛОВОК =====
    add_heading(doc, 'Коммерческое предложение', size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=10)
    add_paragraph(
        doc,
        'на внедрение и сопровождение «Электронной системы Нарядов-Допусков» (далее — Система) '
        'для нужд АО «Каражанбасмунай».',
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    # ===== О РЕШЕНИИ =====
    add_heading(doc, '1. О предлагаемом решении', size=13)
    add_paragraph(
        doc,
        'Готовое к промышленной эксплуатации web-решение, разработанное на современном '
        'технологическом стеке (Python/Django 5, React 19, PostgreSQL). Система полностью '
        'покрывает жизненный цикл наряда-допуска: от создания инициатором — до закрытия '
        'производителем работ и архивирования, с автоматической маршрутизацией согласований, '
        'фиксацией ответственности и поддержкой ЭЦП НУЦ РК.',
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, '2. Ключевые функциональные возможности', size=13)
    add_bullets(doc, [
        'Полный автоматизированный процесс наряда-допуска (создание, согласование, выдача, продление, закрытие).',
        'Гибкая ролевая модель: Администратор, Аудитор, Выдающий наряд, Ответственный руководитель работ, '
        'Допускающий, Производитель работ, Согласующий.',
        'Маршрут согласования по конечному автомату (FSM): Выдающий → Отв. руководитель → Допускающий → '
        'Производитель работ → доп. согласующие (до 5) → основной согласующий (нач. смены / инженер ТБ).',
        'Подписание ЭЦП НУЦ РК (Kalkan, ключи NCALayer / Kaztoken) в веб-интерфейсе; '
        'графическая (рукописная) подпись для исполнителей без ЭЦП.',
        'QR-верификация наряда — подлинность и текущий статус проверяются сканированием с мобильного устройства.',
        'Конструктор шаблонов нарядов: «Огневые работы», «Газоопасные», «Работы на высоте», '
        '«Электротехнические» (с матрицей изоляции и LOTO), «Земляные» и др.',
        'Журнал LOTO (Lock Out / Tag Out): регистрация изоляций, фото-фиксация схем, привязка к нарядам.',
        'Чек-листы безопасности и контроль допусковых процедур.',
        'Автогенерация выходных форм нарядов в DOCX/PDF по утверждённым шаблонам РК.',
        'Двуязычный интерфейс — русский и казахский (полностью локализован).',
        'Дашборды и статистика для роли «Аудитор» (среднее время закрытия, причины отклонений, нагрузка по подразделениям).',
        'Модуль уведомлений: оповещение исполнителей и согласующих о действиях, требующих внимания.',
        'AI-ассистент: контекстная помощь пользователю при заполнении и проверке нарядов.',
        'REST API (DRF) для интеграции с внутренними системами Заказчика (1С, SAP, кадровая система, СЭД).',
        'Готовая мобильная адаптация (PWA) — работа с планшета бригадира на объекте.',
    ])

    add_heading(doc, '3. Состав работ по внедрению', size=13)
    add_bullets(doc, [
        'Установка и настройка программного кода на сервере Заказчика.',
        'Развёртывание СУБД PostgreSQL и обвязки (Nginx, Gunicorn, резервное копирование).',
        'Ввод исходных данных (организационная структура, пользователи, шаблоны нарядов, справочники опасных работ, объекты).',
        'Интеграция с НУЦ РК (NCALayer/Kalkan) для подписания ЭЦП.',
        'Пуско-наладка и проверка работоспособности на тестовом контуре.',
        'Обучение пользователей и администраторов (с выездом на площадку Заказчика).',
        'Передача исходных кодов и эксплуатационной документации.',
        'Гарантийная поддержка в течение 12 месяцев с момента ввода в эксплуатацию.',
    ])

    # ===== ЦЕНА =====
    add_heading(doc, '4. Стоимость', size=13)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    table.autofit = False

    # ----- Заголовок -----
    headers = [
        'Наименование',
        'Стоимость (без НДС),\nтенге',
        'Стоимость (с НДС 16%),\nтенге',
    ]
    for i, h in enumerate(headers):
        c = table.cell(0, i)
        c.text = ''
        set_cell_bg(c, '1F4E79')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)

    # ----- Строка 1 — Внедрение -----
    c = table.cell(1, 0)
    p = c.paragraphs[0]
    r = p.add_run('Внедрение Системы')
    r.bold = True
    r.font.size = Pt(11)
    for sub in [
        '• установка и настройка ПО на сервере Заказчика;',
        '• развёртывание СУБД и инфраструктуры;',
        '• ввод исходных данных и справочников;',
        '• интеграция с НУЦ РК (ЭЦП);',
        '• пуско-наладка и обучение пользователей с выездом;',
        '• передача исходных кодов и документации;',
        '• гарантийная поддержка 12 месяцев — включено.',
    ]:
        sp = c.add_paragraph(sub)
        sp.paragraph_format.left_indent = Cm(0.3)
        for run in sp.runs:
            run.font.size = Pt(10)

    for col in (1, 2):
        cell = table.cell(1, col)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('0')
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = BRAND_BLUE
        p2 = cell.add_paragraph('(бесплатно при заключении договора годового сопровождения)')
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p2.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = GREY

    # ----- Строка 2 — Годовое сопровождение -----
    c = table.cell(2, 0)
    p = c.paragraphs[0]
    r = p.add_run('Обязательное годовое техническое сопровождение')
    r.bold = True
    r.font.size = Pt(11)
    for sub in [
        '• оперативное устранение ошибок и инцидентов;',
        '• мониторинг работоспособности Системы 24/7;',
        '• администрирование ПО и резервное копирование данных;',
        '• обновление ПО при изменении законодательства РК;',
        '• услуги колл-центра в рабочее время Заказчика;',
        '• проверка системы на наличие вредоносного кода;',
        '• выпуск минорных доработок по запросу Заказчика (в рамках договора).',
    ]:
        sp = c.add_paragraph(sub)
        sp.paragraph_format.left_indent = Cm(0.3)
        for run in sp.runs:
            run.font.size = Pt(10)

    price_cols = [
        ('7 000 000 / год', '583 333,33 / месяц'),
        ('8 120 000 / год', '676 666,67 / месяц'),
    ]
    for col, (year_val, month_val) in enumerate(price_cols, start=1):
        cell = table.cell(2, col)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(year_val)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = BRAND_BLUE
        p2 = cell.add_paragraph(month_val)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p2.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = BRAND_BLUE

    # ----- Итог -----
    c = table.cell(3, 0)
    set_cell_bg(c, 'EAF1F8')
    p = c.paragraphs[0]
    r = p.add_run('ИТОГО за первый год эксплуатации')
    r.bold = True
    r.font.size = Pt(12)

    totals = ['7 000 000 тенге', '8 120 000 тенге']
    for col, val in enumerate(totals, start=1):
        cell = table.cell(3, col)
        set_cell_bg(cell, 'EAF1F8')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = BRAND_BLUE

    # Ширина колонок
    for row in table.rows:
        row.cells[0].width = Cm(9)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(4)

    add_paragraph(
        doc,
        'Условия размещения: Система устанавливается и эксплуатируется на серверной '
        'инфраструктуре Заказчика (On-Premise). Все данные нарядов, пользователей и '
        'ЭЦП остаются в периметре АО «Каражанбасмунай».',
        italic=True, size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    # ===== ОПЦИОНАЛЬНО =====
    add_heading(doc, '5. Дополнительно (по опции)', size=13)
    opt = doc.add_table(rows=3, cols=2)
    opt.style = 'Light Grid Accent 1'
    opt.cell(0, 0).text = 'Наименование'
    opt.cell(0, 1).text = 'Стоимость (без НДС), тенге'
    for c in (opt.cell(0, 0), opt.cell(0, 1)):
        set_cell_bg(c, '1F4E79')
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    opt.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    opt.cell(1, 0).text = 'Промышленный планшет Samsung Galaxy Tab Active 5 (1 шт.)'
    opt.cell(1, 1).text = '500 000'
    opt.cell(2, 0).text = 'Kaztoken NFC (носитель ЭЦП, 1 шт.)'
    opt.cell(2, 1).text = '16 000'
    for r in (1, 2):
        opt.cell(r, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== ТРЕБОВАНИЯ К СЕРВЕРУ =====
    add_heading(doc, '6. Требования к серверной инфраструктуре Заказчика', size=13)
    add_paragraph(
        doc,
        'Для промышленной эксплуатации Системы (от 100 до 500 активных пользователей) '
        'рекомендуется следующая конфигурация сервера приложений + СУБД:',
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    srv = doc.add_table(rows=8, cols=3)
    srv.style = 'Light Grid Accent 1'

    headers = ['Параметр', 'Рекомендуется (Prod)', 'Минимум (пилот)']
    for i, h in enumerate(headers):
        c = srv.cell(0, i)
        c.text = ''
        set_cell_bg(c, '1F4E79')
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows_data = [
        ('CPU', 'Intel Xeon / AMD EPYC, 8 vCPU @ 2.4 ГГц+', '4 vCPU @ 2.0 ГГц'),
        ('Оперативная память (RAM)', '32 ГБ DDR4 ECC', '8 ГБ'),
        ('Дисковая подсистема', 'NVMe SSD 1 ТБ (RAID 1/10)', 'SSD 250 ГБ'),
        ('Резервный диск / NAS', 'отдельный том ≥ 1 ТБ для бэкапов', 'внешний диск ≥ 500 ГБ'),
        ('Операционная система', 'Ubuntu Server 22.04 LTS / RHEL 9 / Astra Linux', 'Ubuntu 22.04'),
        ('Сеть', 'внутренняя 1 Гбит/с, доступ к NCALayer/SIGEX', '100 Мбит/с'),
        ('Доп. ПО', 'PostgreSQL 14+, Nginx, Gunicorn, Python 3.11+, Node.js 20+', '— то же —'),
    ]
    for i, (param, prod, mini) in enumerate(rows_data, start=1):
        srv.cell(i, 0).text = param
        srv.cell(i, 1).text = prod
        srv.cell(i, 2).text = mini
        for run in srv.cell(i, 0).paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        for col in (1, 2):
            for run in srv.cell(i, col).paragraphs[0].runs:
                run.font.size = Pt(10)

    add_paragraph(
        doc,
        'Допускается развёртывание в виртуальной среде (VMware vSphere, Proxmox, KVM) '
        'или в частном облаке Заказчика. При необходимости отказоустойчивости — '
        'возможна установка в схеме «сервер приложений + выделенный сервер СУБД» '
        '(конфигурация уточняется на этапе проектирования).',
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10, italic=True, color=GREY,
    )

    # ===== УСЛОВИЯ И СРОКИ =====
    add_heading(doc, '7. Сроки и условия', size=13)
    add_bullets(doc, [
        'Срок внедрения: 4–6 недель с момента подписания договора и предоставления Заказчиком серверных ресурсов.',
        'Гарантийная поддержка: 12 месяцев с момента подписания акта ввода в эксплуатацию — включено в стоимость.',
        'Срок действия настоящего коммерческого предложения: 30 (тридцать) календарных дней с даты подписания.',
        'Форма оплаты: безналичный расчёт, договор и счёт-фактура от [ТОО «___»].',
        'Условие бесплатного внедрения — заключение договора на годовое техническое сопровождение.',
    ])

    # ===== ПОДПИСЬ =====
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, 'С уважением,', size=11)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, '________________________', size=11)
    add_paragraph(doc, '[ФИО руководителя]', size=11, bold=True)
    add_paragraph(doc, 'Директор [ТОО «___»]', size=11, color=GREY)
    add_paragraph(doc, 'Тел.: [+7 (___) ___-__-__]   E-mail: [___@___.kz]', size=10, color=GREY)

    return doc


if __name__ == '__main__':
    out = '/home/sadmin/web/dev/КП_Электронные_наряды-допуски_КБМ.docx'
    make_doc().save(out)
    print(f'OK: {out}')
