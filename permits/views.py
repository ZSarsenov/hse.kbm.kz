# permits/views.py
from django.http import HttpResponse
from django.conf import settings
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
from collections import Counter, defaultdict
from datetime import datetime
from openai import OpenAI
import logging
import time
from rest_framework import viewsets, status, filters
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.views import APIView
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.db import IntegrityError, transaction
from django.db.models import Q
from rest_framework.exceptions import PermissionDenied

from .models import WorkPermit, WorkPermitTemplate, Department, DangerousWorkType, ApprovalStep, Notification
from core.signature import parse_xml_signature_info
from .serializers import (PermitSerializer, WorkPermitTemplateSerializer, DepartamentSerializer,
                          DangerousWorkTypeSerializer, NotificationSerializer)

from .kalkan import Kalkan # временно не используем способ подписания через Kalkan

logger = logging.getLogger(__name__)


def _issuer_admitting_approved_users(permit):
    """Пользователи, уже подписавшие шаги Выдающий и Допускающий (для уведомлений о графической подписи производителя)."""
    users = []
    seen = set()
    for role in ('ISSUER', 'ADMITTING'):
        st = (
            permit.approval_steps.filter(role=role, status='APPROVED')
            .select_related('approver')
            .first()
        )
        if st and st.approver_id and st.approver_id not in seen:
            seen.add(st.approver_id)
            users.append(st.approver)
    return users


def _prior_approved_approvers(permit, before_step_order):
    """Участники с шагами APPROVED до указанного порядка (кто может передать устройство для графической подписи следующего)."""
    users = []
    seen = set()
    for st in (
        permit.approval_steps.filter(
            status='APPROVED',
            step_order__lt=before_step_order,
            approver_id__isnull=False,
        )
        .order_by('step_order')
        .select_related('approver')
    ):
        if st.approver_id not in seen:
            seen.add(st.approver_id)
            users.append(st.approver)
    return users


def _advance_after_approval_step(permit, completed_step, acting_user):
    """Перевод следующего шага в PENDING или финализация наряда; уведомления."""
    next_step = ApprovalStep.objects.filter(
        permit=permit,
        step_order__gt=completed_step.step_order,
    ).order_by('step_order').first()

    if next_step:
        next_step.status = 'PENDING'
        next_step.save()

        if next_step.approver_id:
            Notification.objects.create(
                user=next_step.approver,
                permit_id=permit.id,
                title="Требуется согласование",
                message=f"Наряд №{permit.permit_id} ожидает вашей подписи ({next_step.get_role_display()}).",
            )
        else:
            if next_step.role == 'WORK_PRODUCER':
                for nu in _issuer_admitting_approved_users(permit):
                    Notification.objects.create(
                        user=nu,
                        permit_id=permit.id,
                        title="Графическая подпись производителя",
                        message=(
                            f"Наряд №{permit.permit_id}: внесите графическую подпись производителя работ "
                            f"(исполнитель без ЭЦП) в карточке наряда (вкладка «Описание»)."
                        ),
                    )
            elif next_step.role == 'COORDINATOR':
                sup = (permit.data or {}).get('supervisor') or {}
                if isinstance(sup, dict) and sup.get('external'):
                    for nu in _prior_approved_approvers(permit, next_step.step_order):
                        Notification.objects.create(
                            user=nu,
                            permit_id=permit.id,
                            title="Графическая подпись согласующего",
                            message=(
                                f"Наряд №{permit.permit_id}: внесите графическую подпись согласующего "
                                f"(начальник смены / участка / инженер ТБ без ЭЦП) в карточке наряда (вкладка «Основное»)."
                            ),
                        )

        if permit.initiator_id != acting_user.id:
            ext_prod = completed_step.role == 'WORK_PRODUCER' and completed_step.approver_id is None
            sup_d = (permit.data or {}).get('supervisor') or {}
            ext_coord = (
                completed_step.role == 'COORDINATOR'
                and completed_step.approver_id is None
                and isinstance(sup_d, dict)
                and sup_d.get('external')
            )
            if ext_prod or ext_coord:
                if ext_prod:
                    parts = [
                        f"По наряду №{permit.permit_id} зафиксирована графическая подпись производителя работ "
                        f"({acting_user.get_full_name()})."
                    ]
                else:
                    parts = [
                        f"По наряду №{permit.permit_id} зафиксирована графическая подпись согласующего "
                        f"({acting_user.get_full_name()})."
                    ]
                if next_step.approver_id:
                    parts.append(f"Передан следующему: {next_step.approver.get_full_name()}.")
                Notification.objects.create(
                    user=permit.initiator,
                    permit_id=permit.id,
                    title="Наряд подписан",
                    message=" ".join(parts),
                )
            else:
                if next_step.approver_id:
                    suffix = f" Передан следующему: {next_step.approver.get_full_name()}."
                elif next_step.role == 'WORK_PRODUCER':
                    suffix = " Ожидается графическая подпись производителя работ (исполнитель без ЭЦП)."
                elif next_step.role == 'COORDINATOR':
                    suffix = " Ожидается графическая подпись согласующего (без ЭЦП)."
                else:
                    suffix = ""
                Notification.objects.create(
                    user=permit.initiator,
                    permit_id=permit.id,
                    title="Наряд подписан",
                    message=(
                        f"Пользователь {acting_user.get_full_name()} подписал наряд №{permit.permit_id}."
                        + suffix
                    ),
                )

        if next_step.approver_id:
            print(f"🔔 Уведомление отправлено пользователю {next_step.approver.get_full_name()}")
    else:
        permit.approve_final()
        Notification.objects.create(
            user=permit.initiator,
            permit_id=permit.id,
            title="Наряд утвержден",
            message=f"Ваш наряд №{permit.permit_id} полностью согласован всеми участниками!",
        )
        # Если производитель внешний — уведомляем Выдающего и Допускающего о необходимости закрыть наряд
        prod_data = (permit.data or {}).get('producer') or {}
        if isinstance(prod_data, dict) and prod_data.get('external'):
            for nu in _issuer_admitting_approved_users(permit):
                Notification.objects.create(
                    user=nu,
                    permit_id=permit.id,
                    title="Требуется закрытие наряда",
                    message=(
                        f"Наряд №{permit.permit_id} согласован. Производитель работ выступает без ЭЦП — "
                        f"необходимо закрыть наряд от его имени (кнопка «Закрыть наряд как Производитель работ»)."
                    ),
                )


class WorkPermitTemplateViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Шаблоны нарядов-допусков (только чтение).
    """
    queryset = WorkPermitTemplate.objects.all()
    serializer_class = WorkPermitTemplateSerializer
    permission_classes = [AllowAny]  # на время DEV можно открыть для всех


class WorkPermitViewSet(viewsets.ModelViewSet):
    queryset = WorkPermit.objects.all()
    serializer_class = PermitSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        if user.is_admin or user.is_auditor:
            return WorkPermit.objects.all()
        q = Q(initiator=user) | Q(approval_steps__approver=user)
        if user.username == 'dispatcher_semser':
            q = q | Q(data__notifyFireService=True)
        return WorkPermit.objects.filter(q).distinct()

    @action(detail=False, methods=['get'], url_path='audit_stats')
    def audit_stats(self, request):
        """
        Агрегированная статистика для дашборда аудитора.
        Параметры:
        - date_from=YYYY-MM-DD
        - date_to=YYYY-MM-DD
        - group_by=day|week|month (по умолчанию day)
        """
        user = request.user
        if not (user.is_admin or user.is_auditor):
            raise PermissionDenied("Доступно только для роли AUDITOR/ADMIN.")

        date_from_raw = request.query_params.get('date_from')
        date_to_raw = request.query_params.get('date_to')
        group_by = request.query_params.get('group_by', 'day')
        if group_by not in ('day', 'week', 'month'):
            group_by = 'day'

        date_from = None
        date_to = None
        try:
            if date_from_raw:
                date_from = datetime.strptime(date_from_raw, '%Y-%m-%d').date()
            if date_to_raw:
                date_to = datetime.strptime(date_to_raw, '%Y-%m-%d').date()
        except ValueError:
            return Response(
                {'error': 'Неверный формат даты. Используйте YYYY-MM-DD.'},
                status=status.HTTP_400_BAD_REQUEST
            )

        qs = WorkPermit.objects.select_related('location', 'initiator', 'initiator__department').all()
        if date_from:
            qs = qs.filter(created_at__date__gte=date_from)
        if date_to:
            qs = qs.filter(created_at__date__lte=date_to)

        period_qs = list(
            qs.values(
                'id', 'status', 'created_at', 'location__name', 'data',
                'initiator__department__name'
            )
        )
        total_qs = WorkPermit.objects.all()

        status_order = ['DRAFT', 'PENDING_APPROVAL', 'APPROVED', 'REJECTED', 'CLOSED']
        status_counts = Counter(item['status'] for item in period_qs)
        status_distribution = [
            {'status': code, 'count': status_counts.get(code, 0)}
            for code in status_order
        ]

        trend_counter = defaultdict(int)
        for item in period_qs:
            dt = timezone.localtime(item['created_at'])
            if group_by == 'month':
                key = dt.strftime('%Y-%m')
            elif group_by == 'week':
                iso_year, iso_week, _ = dt.isocalendar()
                key = f'{iso_year}-W{iso_week:02d}'
            else:
                key = dt.strftime('%Y-%m-%d')
            trend_counter[key] += 1
        permits_trend = [{'period': key, 'count': trend_counter[key]} for key in sorted(trend_counter.keys())]

        work_type_counter = Counter()
        location_counter = Counter()
        department_counter = Counter()
        for item in period_qs:
            payload = item.get('data') or {}
            work_type = (payload.get('workName') or payload.get('content') or 'Не указано').strip()
            # "Топ локаций" показывает Участок/Цех (поле department формы) — оно заполняется всегда.
            # Старые поля location__name и workPlace оставлены как fallback для исторических нарядов.
            location_name = (
                payload.get('department')
                or item.get('initiator__department__name')
                or item.get('location__name')
                or payload.get('workPlace')
                or 'Не указано'
            ).strip()
            department_name = (
                item.get('initiator__department__name') or payload.get('department') or 'Не указано'
            ).strip()
            work_type_counter[work_type] += 1
            location_counter[location_name] += 1
            department_counter[department_name] += 1

        top_work_types = [{'name': name, 'count': count} for name, count in work_type_counter.most_common(7)]
        top_locations = [{'name': name, 'count': count} for name, count in location_counter.most_common(7)]
        top_departments = [{'name': name, 'count': count} for name, count in department_counter.most_common(7)]

        total_count = total_qs.count()
        created_in_period = len(period_qs)
        closed_in_period = status_counts.get('CLOSED', 0)
        rejected_in_period = status_counts.get('REJECTED', 0)
        close_rate = round((closed_in_period / created_in_period) * 100, 1) if created_in_period else 0
        reject_rate = round((rejected_in_period / created_in_period) * 100, 1) if created_in_period else 0

        closed_permits = WorkPermit.objects.filter(
            id__in=[item['id'] for item in period_qs],
            status='CLOSED',
            valid_to__isnull=False
        ).only('created_at', 'valid_to')
        close_durations_hours = []
        for permit in closed_permits:
            duration_hours = (permit.valid_to - permit.created_at).total_seconds() / 3600
            if duration_hours >= 0:
                close_durations_hours.append(duration_hours)
        avg_close_time_hours = (
            round(sum(close_durations_hours) / len(close_durations_hours), 1)
            if close_durations_hours else 0
        )

        return Response({
            'filters': {
                'date_from': date_from_raw,
                'date_to': date_to_raw,
                'group_by': group_by,
            },
            'kpi': {
                'total_all_time': total_count,
                'created_in_period': created_in_period,
                'closed_in_period': closed_in_period,
                'rejected_in_period': rejected_in_period,
                'close_rate_percent': close_rate,
                'reject_rate_percent': reject_rate,
                'avg_close_time_hours': avg_close_time_hours,
            },
            'status_distribution': status_distribution,
            'permits_trend': permits_trend,
            'top_work_types': top_work_types,
            'top_locations': top_locations,
            'top_departments': top_departments,
        })

    # API ДЛЯ "МОИ ЗАДАЧИ" (Колокольчик/Меню)
    @action(detail=False, methods=['get'])
    def my_tasks(self, request):
        """
        Возвращает только те наряды, которые СЕЙЧАС ждут подписи текущего юзера.
        """
        user = request.user
        # Ищем шаги, где approver = Я и статус = PENDING
        my_pending_steps = ApprovalStep.objects.filter(approver=user, status='PENDING')

        # Получаем ID нарядов из этих шагов
        permit_ids = my_pending_steps.values_list('permit_id', flat=True)

        permits = WorkPermit.objects.filter(id__in=permit_ids)
        serializer = self.get_serializer(permits, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'], url_path='submit_for_approval')
    def submit(self, request, pk=None):
        """Отправить наряд на согласование (без подписи). Только для черновика или отклонённого."""
        permit = self.get_object()
        if permit.status not in ('DRAFT', 'REJECTED'):
            return Response(
                {'ok': False, 'error': 'Отправить на согласование можно только черновик или отклонённый наряд.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        was_rejected = permit.status == 'REJECTED'
        try:
            # Одна транзакция: при ошибке на любом шаге не остаётся «обрезанной» цепочки (2 из 5 шагов).
            with transaction.atomic():
                permit.submit()
                permit.save()
            # При повторной отправке после отклонения уведомляем того, кто отказал — наряд снова ждёт его подписи
            if was_rejected:
                pending_step = permit.approval_steps.filter(status='PENDING').first()
                if pending_step and pending_step.approver_id and pending_step.approver != permit.initiator:
                    Notification.objects.create(
                        user=pending_step.approver,
                        permit_id=permit.id,
                        title="Наряд повторно отправлен на согласование",
                        message=f"Наряд №{permit.permit_id} исправлен. Требуется ваша подпись повторно ({pending_step.get_role_display()})."
                    )
            # Уведомление противопожарной службе (если включено при создании)
            if permit.data and permit.data.get('notifyFireService'):
                try:
                    from django.contrib.auth import get_user_model
                    User = get_user_model()
                    dispatcher = User.objects.filter(username='dispatcher_semser').first()
                    if dispatcher:
                        work_name = permit.data.get('workName', 'Не указано')
                        Notification.objects.create(
                            user=dispatcher,
                            permit_id=permit.id,
                            title="Уведомление: огневые работы",
                            message=(
                                f"Создан наряд-допуск №{permit.permit_id}.\n"
                                f"Наименование работ: {work_name}.\n"
                                f"Инициатор: {permit.initiator.get_full_name()}."
                            ),
                        )
                        if permit.data.get('callFirePost'):
                            location = permit.data.get('workPlace', 'Не указано')
                            Notification.objects.create(
                                user=dispatcher,
                                permit_id=permit.id,
                                title="🚨 ВЫЗОВ ПОЖАРНОГО ПОСТА",
                                message=(
                                    f"Требуется вызов пожарной бригады!\n"
                                    f"Наряд-допуск №{permit.permit_id}.\n"
                                    f"Наименование работ: {work_name}.\n"
                                    f"Место проведения: {location}.\n"
                                    f"Инициатор: {permit.initiator.get_full_name()}."
                                ),
                            )
                except Exception:
                    pass

            return Response({
                'ok': True,
                'status': f'Наряд №{permit.permit_id} отправлен на согласование.',
                'current_status': permit.status,
                'permit_id': permit.permit_id,
            })
        except IntegrityError as e:
            err = str(e)
            if 'approver_id' in err and 'null' in err.lower():
                return Response(
                    {
                        'ok': False,
                        'error': (
                            'В базе данных для шага согласования всё ещё обязателен согласующий (approver_id NOT NULL). '
                            'Нужна миграция permits 0016 на том же сервере и той же БД, что использует это приложение: '
                            'python manage.py migrate permits'
                        ),
                    },
                    status=status.HTTP_400_BAD_REQUEST,
                )
            return Response({'ok': False, 'error': err}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({'ok': False, 'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    # --- МЕТОД ПОДПИСАНИЯ ---
    @action(detail=True, methods=["post"], url_path="sign", permission_classes=[IsAuthenticated])
    def sign(self, request, pk=None):
        permit = self.get_object()
        user = request.user
        signed_xml = request.data.get('signed_xml')

        if not signed_xml:
            return Response({"ok": False, "error": "Нет данных подписи"}, status=400)

        # Подписание только для наряда уже на согласовании (черновик отправляют через submit_for_approval)
        if permit.status == 'DRAFT':
            return Response(
                {"ok": False, "error": "Сначала нажмите «Отправить на согласование» на странице наряда."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # ---------------------------------------------------------------
        # ПОИСК ТЕКУЩЕГО ШАГА (с поддержкой нескольких ролей)
        # ---------------------------------------------------------------
        requested_role = request.data.get('role')  # Параметр роли из запроса
        
        # Ищем все PENDING шаги пользователя
        pending_steps = ApprovalStep.objects.filter(
            permit=permit,
            approver=user,
            status='PENDING'
        )
        
        if not pending_steps.exists():
            return Response(
                {"ok": False, "error": "Вы не можете подписать этот наряд (нет активного шага)."},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Если указана роль - ищем конкретный шаг
        if requested_role:
            try:
                current_step = pending_steps.get(role=requested_role)
            except ApprovalStep.DoesNotExist:
                available_roles = [s.get_role_display() for s in pending_steps]
                return Response(
                    {
                        "ok": False,
                        "error": f"Роль '{requested_role}' не найдена среди ваших активных шагов.",
                        "available_roles": [s.role for s in pending_steps],
                        "available_roles_display": available_roles
                    },
                    status=status.HTTP_400_BAD_REQUEST
                )
        else:
            # Если роль не указана - проверяем количество
            if pending_steps.count() == 1:
                current_step = pending_steps.first()
            else:
                # Несколько шагов - требуем указать роль
                available_roles = [{"role": s.role, "role_display": s.get_role_display(), "step_order": s.step_order} 
                                 for s in pending_steps.order_by('step_order')]
                return Response(
                    {
                        "ok": False,
                        "error": "У вас несколько активных ролей для подписания. Укажите параметр 'role'.",
                        "available_roles": available_roles
                    },
                    status=status.HTTP_400_BAD_REQUEST
                )

        # ---------------------------------------------------------------
        # 3. ПРОВЕРКА ЭЦП
        # ---------------------------------------------------------------
        try:
            cert_info = parse_xml_signature_info(signed_xml)
        except Exception as e:
            return Response({"ok": False, "error": f"Ошибка чтения ЭЦП: {str(e)}"}, status=400)

        # Проверка ИИН
        sign_iin = cert_info.get('iin')
        if not sign_iin or sign_iin != user.iin:
            return Response(
                {"ok": False, "error": f"ИИН в ЭЦП ({sign_iin}) не совпадает с вашим ({user.iin})."},
                status=400
            )

        # Проверка БИН
        sign_bin = cert_info.get('bin')
        user_bin = user.bin
        target_bin = user_bin if user_bin else '950540000524'

        if not sign_bin:
            return Response({"ok": False, "error": "Нужна ЭЦП юридического лица (GOST) с БИН."}, status=400)

        if sign_bin != target_bin:
            return Response({"ok": False, "error": f"БИН организации не совпадает ({sign_bin} != {target_bin})."},
                            status=400)

        # ---------------------------------------------------------------
        # 4. СОХРАНЕНИЕ ПОДПИСИ ТЕКУЩЕГО ШАГА
        # ---------------------------------------------------------------
        current_step.status = 'APPROVED'
        current_step.signed_xml = signed_xml
        # Сохраняем дату текстом для истории
        current_step.signed_at = timezone.now()
        current_step.signer_details = {
            "iin": cert_info.get('iin'),
            "bin": cert_info.get('bin'),
            "fio": cert_info.get('subject'), # Или распарсить CN из subject
            "org_name": cert_info.get('org_name'),
            "date": timezone.now().isoformat()
        }
        current_step.save()

        # ---------------------------------------------------------------
        # 5. ПЕРЕХОД ХОДА ИЛИ ФИНАЛИЗАЦИЯ
        # ---------------------------------------------------------------
        _advance_after_approval_step(permit, current_step, user)

        permit.save()

        next_for_response = ApprovalStep.objects.filter(
            permit=permit,
            step_order__gt=current_step.step_order,
        ).order_by('step_order').first()
        return Response({
            "ok": True,
            "status": "Наряд успешно подписан",
            "next_step": next_for_response.role if next_for_response else "Согласование завершено",
        })

    # --- МЕТОД ОТКЛОНЕНИЯ ---
    @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
    def reject(self, request, pk=None):
        """
        Отклонение наряда текущим согласующим.
        """
        permit = self.get_object()
        user = request.user
        reason = request.data.get('reason')

        if not reason:
            return Response({"ok": False, "error": "Необходимо указать причину отклонения."}, status=400)

        # 1. Ищем шаг, который сейчас ждет подписи от этого пользователя
        try:
            step = ApprovalStep.objects.get(permit=permit, approver=user, status='PENDING')
        except ApprovalStep.DoesNotExist:
            return Response(
                {"ok": False, "error": "Вы не можете отклонить этот наряд (нет активного шага согласования)."},
                status=403
            )

        # 2. Отмечаем шаг как ОТКЛОНЕННЫЙ
        step.status = 'REJECTED'
        step.signed_at = timezone.now()
        # СОХРАНЯЕМ ПРИЧИНУ ОТКАЗА
        step.rejection_reason = reason
        step.save()

        # 3. Переводим весь наряд в статус REJECTED
        try:
            permit.reject()  # Вызывает transition из models.py
            permit.save()
        except Exception as e:
            return Response({"ok": False, "error": str(e)}, status=400)

        # 4. Уведомляем создателя
        Notification.objects.create(
            user=permit.initiator,
            permit_id=permit.id,
            title="Наряд отклонен ❌",
            message=f"Пользователь {user.get_full_name()} отклонил ваш наряд. Причина: {reason}"
        )

        return Response({"ok": True, "status": "Наряд отклонен"})

    @action(detail=False, methods=['get'])
    def my_tasks(self, request):
        user = request.user
        # Находим шаги, где я - approver И статус - PENDING
        my_steps = ApprovalStep.objects.filter(approver=user, status='PENDING')
        permit_ids = my_steps.values_list('permit_id', flat=True)
        permits = WorkPermit.objects.filter(id__in=permit_ids)
        serializer = self.get_serializer(permits, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
    def duplicate(self, request, pk=None):
        """
        Создание копии наряда (черновика) на основе существующего.
        """
        try:
            original = self.get_object()

            # 1. Копируем JSON данные
            # Если data нет, используем пустой dict, чтобы не упало
            new_data = (original.data or {}).copy()

            # 2. Очищаем данные (подписи бригады не копируем — новый наряд подписывают заново)
            new_data['dateStart'] = ""
            new_data['dateEnd'] = ""
            new_data['riskApprovedBy'] = ""
            new_data['extensions'] = []
            new_data['brigade_signatures'] = []
            new_data.pop('producer_signature', None)
            new_data.pop('supervisor_signature', None)

            # 3. Генерируем номер по стандарту: OR-2026-00001
            current_year = time.strftime('%Y')
            prefix = f"OR-{current_year}-"
            last_permit = WorkPermit.objects.filter(
                permit_id__startswith=prefix
            ).order_by('-permit_id').first()
            if last_permit:
                last_number = int(last_permit.permit_id.split('-')[-1])
                next_number = last_number + 1
            else:
                next_number = 1
            new_permit_id = f"{prefix}{next_number:05d}"

            # 4. Создаем новый объект
            new_permit = WorkPermit.objects.create(
                initiator=request.user,
                permit_id=new_permit_id,  # 👈 ОБЯЗАТЕЛЬНОЕ ПОЛЕ
                status='DRAFT',
                template=original.template,  # Не забываем сам шаблон (объект)

                # category=original.category,
                location=original.location,  # Копируем объект локации

                data=new_data
            )

            return Response({"ok": True, "id": new_permit.id, "message": "Наряд скопирован"})

        except Exception as e:
            print(f"Ошибка при дублировании: {e}")  # Выведет ошибку в консоль сервера
            return Response({"ok": False, "error": str(e)}, status=400)

    @action(detail=False, methods=['get'], url_path='export_journal')
    def export_journal(self, request):
        """
        Экспорт журнала (закрытые/отклонённые наряды) в xlsx.
        Параметры: date_from, date_to (YYYY-MM-DD), status (CLOSED|REJECTED) — необязательные.
        Если записей нет — файл скачивается пустым (только заголовки).
        """
        from django.http import HttpResponse
        from openpyxl import Workbook
        from openpyxl.utils import get_column_letter
        from datetime import datetime
        from io import BytesIO

        qs = self.get_queryset().filter(status__in=['CLOSED', 'REJECTED']).order_by('-valid_from', '-created_at')

        date_from_s = request.query_params.get('date_from')
        date_to_s = request.query_params.get('date_to')
        status_param = request.query_params.get('status')
        if status_param in ('CLOSED', 'REJECTED'):
            qs = qs.filter(status=status_param)
        if date_from_s:
            try:
                dt_from = datetime.strptime(date_from_s, '%Y-%m-%d')
                qs = qs.filter(valid_from__date__gte=dt_from.date())
            except ValueError:
                pass
        if date_to_s:
            try:
                dt_to = datetime.strptime(date_to_s, '%Y-%m-%d')
                qs = qs.filter(valid_from__date__lte=dt_to.date())
            except ValueError:
                pass

        wb = Workbook()
        ws = wb.active
        ws.title = "Журнал нарядов-допусков"
        headers = [
            '№', 'Начало (Первичный допуск)', 'Окончание (Первичный допуск)', '№ наряда-допуска',
            'Лицо, выдавшее наряд', 'Наименование работ', 'Статус'
        ]
        for col, h in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=h)
        for idx, permit in enumerate(qs, 1):
            issuer_step = permit.approval_steps.filter(role='ISSUER').first()
            issuer_name = issuer_step.approver.get_full_name() if issuer_step and issuer_step.approver else (permit.initiator.get_full_name() if permit.initiator else '—')
            valid_from_str = permit.valid_from.strftime('%d.%m.%Y %H:%M') if permit.valid_from else '—'
            valid_to_str = permit.valid_to.strftime('%d.%m.%Y %H:%M') if permit.valid_to else '—'
            work_name = (permit.data or {}).get('workName') or '—'
            status_display = dict(WorkPermit.STATUS_CHOICES).get(permit.status, permit.status)
            ws.append([idx, valid_from_str, valid_to_str, permit.permit_id, issuer_name, work_name, status_display])
        for col in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 18
        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        response = HttpResponse(buf.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        fname = f"journal_{date_from_s or 'all'}_{date_to_s or 'all'}_{status_param or 'all'}.xlsx"
        response['Content-Disposition'] = f'attachment; filename="{fname}"'
        return response

    @action(detail=True, methods=['post'], url_path='brigade_signature')
    def brigade_signature(self, request, pk=None):
        """
        Сохранение подписи члена бригады (рисование на экране).
        Только для наряда в статусе Согласован. Тело: member_index (int), signature (image file).
        """
        permit = self.get_object()
        if permit.status != 'APPROVED':
            return Response(
                {'error': 'Подписи бригады принимаются только для согласованного наряда.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        member_index = request.data.get('member_index')
        if member_index is None:
            return Response({'error': 'Укажите member_index (номер члена бригады, с 0).'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            member_index = int(member_index)
        except (TypeError, ValueError):
            return Response({'error': 'member_index должен быть числом.'}, status=status.HTTP_400_BAD_REQUEST)
        team = permit.data.get('teamMembers') or []
        if member_index < 0 or member_index >= len(team):
            return Response({'error': 'Недопустимый номер члена бригады.'}, status=status.HTTP_400_BAD_REQUEST)
        image_file = request.FILES.get('signature')
        if not image_file:
            return Response({'error': 'Приложите файл подписи (signature).'}, status=status.HTTP_400_BAD_REQUEST)
        ct = (image_file.content_type or '').lower()
        if ct and not ct.startswith('image/'):
            return Response({'error': f'Разрешены только изображения (получен {image_file.content_type}).'}, status=status.HTTP_400_BAD_REQUEST)
        if image_file.size > 2 * 1024 * 1024:
            return Response({'error': 'Размер файла не более 2 МБ.'}, status=status.HTTP_400_BAD_REQUEST)
        import os
        try:
            rel_dir = os.path.join('brigade_signatures', str(permit.pk))
            dest_dir = os.path.join(settings.MEDIA_ROOT, rel_dir)
            os.makedirs(dest_dir, exist_ok=True)
            fname = f'sign_{member_index}.png'
            rel_path = os.path.join(rel_dir, fname).replace('\\', '/')
            full_path = os.path.join(settings.MEDIA_ROOT, rel_dir, fname)
            with open(full_path, 'wb') as f:
                for chunk in image_file.chunks():
                    f.write(chunk)
            data = dict(permit.data) if permit.data else {}
            sigs = data.get('brigade_signatures')
            if sigs is None:
                sigs = [None] * len(team)
            elif isinstance(sigs, dict):
                sigs = [sigs.get(str(i)) for i in range(len(team))]
            while len(sigs) <= member_index:
                sigs.append(None)
            sigs[member_index] = rel_path
            data['brigade_signatures'] = sigs
            team = data.get('teamMembers') or []
            # Дата/время инструктажа всегда фиксируется по моменту подписи члена бригады
            if member_index < len(team):
                team[member_index]['instructedAt'] = timezone.now().isoformat()
                data['teamMembers'] = team
            permit.data = data
            permit.save(update_fields=['data'])
            return Response({'ok': True, 'member_index': member_index, 'signature_path': rel_path})
        except Exception as e:
            return Response(
                {'error': f'Ошибка сохранения подписи: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['post'], url_path='add_brigade_member')
    def add_brigade_member(self, request, pk=None):
        """Добавление члена бригады к согласованному наряду."""
        permit = self.get_object()
        if permit.status != 'APPROVED':
            return Response({'error': 'Добавить члена бригады можно только к согласованному наряду.'}, status=400)

        name = (request.data.get('name') or '').strip()
        role = (request.data.get('role') or '').strip()
        if not name:
            return Response({'error': 'Укажите ФИО члена бригады.'}, status=400)

        admitting_step = permit.approval_steps.filter(role='ADMITTING').first()
        admitting_name = ''
        if admitting_step and admitting_step.approver:
            admitting_name = admitting_step.approver.get_full_name()
        elif permit.data:
            adm = permit.data.get('admitting') or {}
            if isinstance(adm, dict):
                admitting_name = adm.get('name', '')

        data = dict(permit.data) if permit.data else {}
        team = data.get('teamMembers') or []
        team.append({
            'name': name,
            'role': role,
            'instructedBy': admitting_name,
            # instructedAt НЕ заполняем при добавлении — дата ставится по моменту подписи члена бригады
            'instructedAt': '',
        })
        data['teamMembers'] = team
        permit.data = data
        permit.save(update_fields=['data'])

        return Response({'ok': True, 'member_index': len(team) - 1, 'total': len(team)})

    @action(detail=True, methods=['post'], url_path='producer_signature')
    def producer_signature(self, request, pk=None):
        """
        Графическая подпись производителя работ, если исполнитель без ЭЦП (ввод ФИО вручную).
        Право внести: только Выдающий или Допускающий после своей ЭЦП-подписи.
        """
        permit = self.get_object()
        if permit.status != 'PENDING_APPROVAL':
            return Response(
                {'error': 'Подпись производителя доступна только на этапе согласования.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        producer_step = permit.approval_steps.filter(
            role='WORK_PRODUCER', status='PENDING', approver__isnull=True,
        ).first()
        if not producer_step:
            return Response(
                {'error': 'Нет активного шага «Производитель работ» для исполнителя без ЭЦП.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        prod = permit.data.get('producer') if permit.data else {}
        if not (isinstance(prod, dict) and prod.get('external')):
            return Response(
                {'error': 'Производитель работ указан из базы — используйте ЭЦП.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        user = request.user
        issuer_ok = permit.approval_steps.filter(
            role='ISSUER', approver=user, status='APPROVED',
        ).exists()
        admit_ok = permit.approval_steps.filter(
            role='ADMITTING', approver=user, status='APPROVED',
        ).exists()
        if not (issuer_ok or admit_ok):
            return Response(
                {'error': 'Графическую подпись производителя могут внести только Выдающий или Допускающий после своей подписи.'},
                status=status.HTTP_403_FORBIDDEN,
            )

        image_file = request.FILES.get('signature')
        if not image_file:
            return Response({'error': 'Приложите файл подписи (signature).'}, status=status.HTTP_400_BAD_REQUEST)
        ct = (image_file.content_type or '').lower()
        if ct and not ct.startswith('image/'):
            return Response({'error': f'Разрешены только изображения (получен {image_file.content_type}).'}, status=status.HTTP_400_BAD_REQUEST)
        if image_file.size > 2 * 1024 * 1024:
            return Response({'error': 'Размер файла не более 2 МБ.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            rel_dir = os.path.join('brigade_signatures', str(permit.pk))
            dest_dir = os.path.join(settings.MEDIA_ROOT, rel_dir)
            os.makedirs(dest_dir, exist_ok=True)
            fname = 'producer.png'
            rel_path = os.path.join(rel_dir, fname).replace('\\', '/')
            full_path = os.path.join(settings.MEDIA_ROOT, rel_dir, fname)
            with open(full_path, 'wb') as f:
                for chunk in image_file.chunks():
                    f.write(chunk)

            data = dict(permit.data) if permit.data else {}
            data['producer_signature'] = rel_path
            permit.data = data

            producer_step.status = 'APPROVED'
            producer_step.signed_at = timezone.now()
            producer_step.signer_details = {
                'graphic': True,
                'recorded_by': user.id,
                'recorded_by_name': user.get_full_name(),
            }
            producer_step.save(update_fields=['status', 'signed_at', 'signer_details'])

            permit.save(update_fields=['data'])
            _advance_after_approval_step(permit, producer_step, user)
            permit.save()

            return Response({'ok': True, 'signature_path': rel_path})
        except Exception as e:
            return Response(
                {'error': f'Ошибка сохранения подписи: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    @action(detail=True, methods=['post'], url_path='supervisor_signature')
    def supervisor_signature(self, request, pk=None):
        """
        Графическая подпись основного согласующего (Нач. смены / участка / инженер ТБ), если указан без ЭЦП.
        Внести могут пользователи, уже подписавшие любой шаг до этого согласующего (передать планшет/телефон).
        """
        permit = self.get_object()
        if permit.status != 'PENDING_APPROVAL':
            return Response(
                {'error': 'Подпись согласующего доступна только на этапе согласования.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        supervisor_step = permit.approval_steps.filter(
            role='COORDINATOR', status='PENDING', approver__isnull=True,
        ).order_by('-step_order').first()
        if not supervisor_step:
            return Response(
                {'error': 'Нет активного шага основного согласующего без ЭЦП.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        sup = permit.data.get('supervisor') if permit.data else {}
        if not (isinstance(sup, dict) and sup.get('external')):
            return Response(
                {'error': 'Согласующий указан из базы — используйте ЭЦП.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        user = request.user
        prior_ids = set(
            permit.approval_steps.filter(
                status='APPROVED',
                step_order__lt=supervisor_step.step_order,
                approver_id__isnull=False,
            ).values_list('approver_id', flat=True)
        )
        if not getattr(user, 'is_admin', False) and user.id not in prior_ids:
            return Response(
                {
                    'error': (
                        'Графическую подпись могут внести только участники, уже подписавшие шаги '
                        'до этого согласующего (или администратор).'
                    )
                },
                status=status.HTTP_403_FORBIDDEN,
            )

        image_file = request.FILES.get('signature')
        if not image_file:
            return Response({'error': 'Приложите файл подписи (signature).'}, status=status.HTTP_400_BAD_REQUEST)
        ct = (image_file.content_type or '').lower()
        if ct and not ct.startswith('image/'):
            return Response({'error': f'Разрешены только изображения (получен {image_file.content_type}).'}, status=status.HTTP_400_BAD_REQUEST)
        if image_file.size > 2 * 1024 * 1024:
            return Response({'error': 'Размер файла не более 2 МБ.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            rel_dir = os.path.join('brigade_signatures', str(permit.pk))
            dest_dir = os.path.join(settings.MEDIA_ROOT, rel_dir)
            os.makedirs(dest_dir, exist_ok=True)
            fname = 'supervisor.png'
            rel_path = os.path.join(rel_dir, fname).replace('\\', '/')
            full_path = os.path.join(settings.MEDIA_ROOT, rel_dir, fname)
            with open(full_path, 'wb') as f:
                for chunk in image_file.chunks():
                    f.write(chunk)

            data = dict(permit.data) if permit.data else {}
            data['supervisor_signature'] = rel_path
            permit.data = data

            supervisor_step.status = 'APPROVED'
            supervisor_step.signed_at = timezone.now()
            supervisor_step.signer_details = {
                'graphic': True,
                'recorded_by': user.id,
                'recorded_by_name': user.get_full_name(),
            }
            supervisor_step.save(update_fields=['status', 'signed_at', 'signer_details'])

            permit.save(update_fields=['data'])
            _advance_after_approval_step(permit, supervisor_step, user)
            permit.save()

            return Response({'ok': True, 'signature_path': rel_path})
        except Exception as e:
            return Response(
                {'error': f'Ошибка сохранения подписи: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    # --- МЕТОД СКАЧИВАНИЯ ---
    # 👇 НОВЫЙ МЕТОД ГЕНЕРАЦИИ PDF
    @action(detail=True, methods=['get'])
    def download_pdf(self, request, pk=None):
        """
        Генерация PDF с использованием шрифта Times New Roman (кириллица).
        """
        # Импорты лучше держать в начале файла, но можно и тут
        from django.http import HttpResponse
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os

        try:
            permit = self.get_object()

            response = HttpResponse(content_type='application/pdf')
            filename = f"Permit_{permit.pk}.pdf"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'

            p = canvas.Canvas(response, pagesize=A4)
            width, height = A4

            # --- 1. НАСТРОЙКА ШРИФТА (Times New Roman) ---

            # Указываем точный путь к файлу, который вы загрузили
            font_path = '/home/sadmin/web/prod/static/fonts/times.ttf'  # 👈 ИЗМЕНИЛИ ПУТЬ

            # Имя, под которым шрифт будет известен внутри PDF
            font_name = 'TimesNewRoman'  # 👈 ДАЛИ ИМЯ

            if os.path.exists(font_path):
                try:
                    # Регистрируем шрифт
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    current_font = font_name
                except Exception as e:
                    print(f"Ошибка регистрации шрифта: {e}")
                    current_font = "Helvetica"
            else:
                print(f"⚠️ Файл шрифта не найден по пути: {font_path}")
                current_font = "Helvetica"  # Если файл не найден, будут квадраты, но не ошибка

            # --- 2. РИСОВАНИЕ ---

            y = height - 50
            x = 50

            # Заголовок (Жирный шрифт, размер 16)
            p.setFont(current_font, 16)
            p.drawString(x, y, f"Наряд-допуск № {permit.pk}")
            y -= 30

            # Основной текст (Обычный шрифт, размер 12)
            p.setFont(current_font, 12)

            def draw_line(label, value):
                nonlocal y
                # Рисуем строку
                p.drawString(x, y, f"{label}: {value}")
                y -= 20

            # Статус
            draw_line("Статус", permit.get_status_display())

            # Инициатор
            initiator = permit.initiator.get_full_name() if permit.initiator else "Неизвестно"
            draw_line("Инициатор", initiator)

            # Место работ
            loc_name = permit.location.name if permit.location else "Не указано"
            draw_line("Место работ", loc_name)

            # Период
            valid_from = permit.valid_from.strftime('%d.%m.%Y %H:%M') if permit.valid_from else "..."
            valid_to = permit.valid_to.strftime('%d.%m.%Y %H:%M') if permit.valid_to else "..."
            draw_line("Период", f"{valid_from} - {valid_to}")

            y -= 10
            p.drawString(x, y, "Описание работ:")
            y -= 20

            # Описание
            content = "Нет описания"
            if permit.data and 'content' in permit.data:
                content = permit.data['content']

            # Простая обрезка длинного текста (чтобы не улетел за край)
            # В будущем заменим на перенос строк
            max_len = 75
            if len(content) > max_len:
                p.drawString(x, y, content[:max_len])
                y -= 15
                p.drawString(x, y, content[max_len:max_len * 2] + "...")
            else:
                p.drawString(x, y, content)

            y -= 40

            # --- 3. ЛИСТ СОГЛАСОВАНИЯ ---
            p.setFont(current_font, 14)
            p.drawString(x, y, "Лист согласования:")
            y -= 25

            p.setFont(current_font, 10)  # Шрифт поменьше для списка

            steps = permit.approval_steps.all().order_by('step_order')
            for step in steps:
                role = step.get_role_display()
                approver = step.approver.get_full_name() if step.approver else "—"

                status_label = "ОЖИДАНИЕ"
                if step.status == 'APPROVED':
                    status_label = "ПОДПИСАНО"
                elif step.status == 'REJECTED':
                    status_label = "ОТКАЗАНО"
                elif step.status == 'PENDING':
                    status_label = "НА ПОДПИСАНИИ"

                date_str = ""
                if step.signed_at:
                    date_str = step.signed_at.strftime('%d.%m.%Y %H:%M')

                line = f"{step.step_order}. {role}: {approver} — {status_label} {date_str}"

                # Если место на странице кончилось
                if y < 50:
                    p.showPage()  # Новая страница
                    p.setFont(current_font, 10)
                    y = height - 50

                p.drawString(x, y, line)
                y -= 15

            p.showPage()
            p.save()
            return response

        except Exception as e:
            print(f"🔥 PDF ERROR: {e}")
            return HttpResponse(f"Ошибка генерации PDF: {e}", status=500)

    def _get_rendered_doc(self, permit, qr_url):
        """Собирает и рендерит DOCX по наряду, возвращает DocxTemplate или None."""
        from django.conf import settings
        from docxtpl import DocxTemplate, InlineImage
        from docx.shared import Mm
        import os
        import qrcode
        from io import BytesIO
        from zoneinfo import ZoneInfo
        from datetime import datetime

        kz_tz = ZoneInfo('Asia/Almaty')
        template_path = os.path.join(settings.BASE_DIR, 'templates', 'docx', 'dangerous_permits_rus.docx')

        if not os.path.exists(template_path):
            return None

        doc = DocxTemplate(template_path)

        # --- QR CODE (ссылка на верификацию — скачивание документа без авторизации) ---
        qr = qrcode.QRCode(box_size=10, border=1)
        qr.add_data(qr_url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        img_buffer = BytesIO()
        img.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        qr_image = InlineImage(doc, img_buffer, width=Mm(25))

        # --- ФУНКЦИИ ДАТ ---
        def format_date(dt):
            if not dt: return '"___" _________ 20___г.'
            return dt.astimezone(kz_tz).strftime('%d.%m.%Y %H:%M')

        def format_str_date(date_str):
            if not date_str: return ""
            try:
                dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                return dt.astimezone(kz_tz).strftime('%d.%m.%Y %H:%M')
            except Exception:
                return date_str

        def get_person(role_enum, json_key):
            step = permit.approval_steps.filter(role=role_enum).first()
            data = {'name': '_________________', 'job': '_________________', 'date': '"___" _________ 20___г.'}
            if step and step.approver:
                user = step.approver
                data['name'] = user.get_full_name()
                data['job'] = getattr(user, 'job_title', getattr(user, 'position', 'Должность не указана'))
                if step.status == 'APPROVED' and step.signed_at:
                    data['date'] = format_date(step.signed_at)
                return data
            # Производитель без учётной записи (исполнитель без ЭЦП): шаг без approver, ФИО/должность в data.producer
            if step and role_enum == 'WORK_PRODUCER' and step.approver_id is None and permit.data:
                p_data = permit.data.get(json_key) or {}
                if isinstance(p_data, dict) and p_data.get('external'):
                    line = (p_data.get('name') or p_data.get('freeText') or '').strip()
                    data['name'] = line or '_________________'
                    has_graphic_sig = bool(permit.data.get('producer_signature'))
                    data['job'] = '' if has_graphic_sig else '_________________'
                    if step.status == 'APPROVED' and step.signed_at:
                        data['date'] = format_date(step.signed_at)
                    return data
            if step and role_enum == 'COORDINATOR' and step.approver_id is None and permit.data:
                s_data = permit.data.get(json_key) or {}
                if isinstance(s_data, dict) and s_data.get('external'):
                    line = (s_data.get('name') or s_data.get('freeText') or '').strip()
                    data['name'] = line or '_________________'
                    has_graphic_sig = bool(permit.data.get('supervisor_signature'))
                    data['job'] = '' if has_graphic_sig else '_________________'
                    if step.status == 'APPROVED' and step.signed_at:
                        data['date'] = format_date(step.signed_at)
                    return data
            if permit.data and json_key in permit.data:
                p_data = permit.data[json_key]
                if isinstance(p_data, dict) and p_data.get('name'):
                    data['name'] = p_data.get('name')
                    data['job'] = p_data.get('position', '_________________')
                    if permit.created_at:
                        data['date'] = format_date(permit.created_at)
            return data

        prod_info = get_person('WORK_PRODUCER', 'producer')
        admit_info = get_person('ADMITTING', 'admitting')
        resp_info = get_person('RESPONSIBLE', 'responsible')

        # Основной согласующий — последний COORDINATOR по step_order
        coord_steps = list(permit.approval_steps.filter(role='COORDINATOR').order_by('step_order'))
        empty_person = {'name': '_________________', 'job': '_________________', 'date': '"___" _________ 20___г.'}
        if coord_steps:
            main_coord_step = coord_steps[-1]
            user = main_coord_step.approver
            if user:
                super_info = {
                    'name': user.get_full_name(),
                    'job': getattr(user, 'job_title', getattr(user, 'position', 'Должность не указана')),
                    'date': format_date(main_coord_step.signed_at) if main_coord_step.status == 'APPROVED' and main_coord_step.signed_at else '"___" _________ 20___г.',
                }
            elif main_coord_step.approver_id is None and permit.data:
                s_data = permit.data.get('supervisor') or {}
                if isinstance(s_data, dict) and s_data.get('external'):
                    line = (s_data.get('name') or s_data.get('freeText') or '').strip()
                    super_info = {
                        'name': line or '_________________',
                        'job': '' if bool(permit.data.get('supervisor_signature')) else '_________________',
                        'date': format_date(main_coord_step.signed_at) if main_coord_step.status == 'APPROVED' and main_coord_step.signed_at else '"___" _________ 20___г.',
                    }
                else:
                    super_info = empty_person.copy()
            else:
                super_info = empty_person.copy()
        else:
            super_info = get_person('COORDINATOR', 'supervisor')
            if '___' in super_info['name']:
                super_info = get_person('SHIFT_SUPERVISOR', 'supervisor')

        # Дополнительные согласующие (все COORDINATOR кроме последнего) — секции 8.1-8.5
        additional_coords_list = []
        if len(coord_steps) > 1:
            for step in coord_steps[:-1]:
                user = step.approver
                if user:
                    additional_coords_list.append({
                        'name': user.get_full_name(),
                        'job': getattr(user, 'job_title', getattr(user, 'position', 'Должность не указана')),
                        'date': format_date(step.signed_at) if step.status == 'APPROVED' and step.signed_at else '"___" _________ 20___г.',
                    })

        issuer_info = get_person('ISSUER', 'issuer')
        if '___' in issuer_info['name']:
            issuer_info['name'] = permit.initiator.get_full_name()
            issuer_info['job'] = getattr(permit.initiator, 'job_title', '')
            issuer_info['date'] = format_date(permit.created_at)

        raw_team = permit.data.get('teamMembers', [])
        sig_paths = permit.data.get('brigade_signatures') or []
        if isinstance(sig_paths, dict):
            sig_paths = [sig_paths.get(str(i)) for i in range(len(raw_team))]
        brigade_list = []
        for idx, m in enumerate(raw_team, start=1):
            row = {
                'i': idx,
                'date': format_str_date(m.get('instructedAt')),
                'name': m.get('name', ''),
                'job': m.get('role', ''),
                'instr_by': admit_info['name'],
            }
            sig_path = sig_paths[idx - 1] if idx - 1 < len(sig_paths) else None
            if sig_path:
                full_path = os.path.join(settings.MEDIA_ROOT, sig_path)
                if os.path.isfile(full_path):
                    try:
                        with open(full_path, 'rb') as f:
                            row['signature_img'] = InlineImage(doc, BytesIO(f.read()), width=Mm(20))
                    except Exception:
                        row['signature_img'] = None
                else:
                    row['signature_img'] = None
            else:
                row['signature_img'] = None
            brigade_list.append(row)
        team_count = len(brigade_list)

        raw_ext = permit.data.get('extensions', [])
        extension_list = [
            {
                'date': format_str_date(ext.get('dateTime')),
                'prod_out': ext.get('producerHandOverName', ''),
                'count': ext.get('incomingTeamCount', ''),
                'prod_in': ext.get('producerTakeOverName', ''),
                'admin': ext.get('admittingName', ''),
            }
            for ext in raw_ext
        ]

        # Графическая подпись производителя (исполнитель без ЭЦП)
        producer_sig_img = ''
        producer_close_sig_img = ''
        supervisor_sig_img = ''
        if permit.data:
            for sig_key, var_name in [
                ('producer_signature', 'producer_sig_img'),
                ('producer_close_signature', 'producer_close_sig_img'),
                ('supervisor_signature', 'supervisor_sig_img'),
            ]:
                sig_rel = permit.data.get(sig_key)
                if sig_rel:
                    sig_full = os.path.join(settings.MEDIA_ROOT, sig_rel)
                    if os.path.isfile(sig_full):
                        try:
                            with open(sig_full, 'rb') as f:
                                img = InlineImage(doc, BytesIO(f.read()), height=Mm(7))
                            if var_name == 'producer_sig_img':
                                producer_sig_img = img
                            elif var_name == 'producer_close_sig_img':
                                producer_close_sig_img = img
                            else:
                                supervisor_sig_img = img
                        except Exception:
                            pass

        context = {
            'qr_code': qr_image,
            'permit_id': permit.permit_id,
            'department': permit.data.get('department', '') or (permit.location.department.name if (permit.location and permit.location.department) else "Не указано"),
            'producer_name': prod_info['name'], 'producer_job': prod_info['job'], 'producer_date': prod_info['date'],
            'producer_signature_img': producer_sig_img,
            'producer_close_signature_img': producer_close_sig_img,
            'supervisor_signature_img': supervisor_sig_img,
            'admitting_name': admit_info['name'], 'admitting_job': admit_info['job'], 'admitting_date': admit_info['date'],
            'responsible_name': resp_info['name'], 'responsible_job': resp_info['job'], 'responsible_date': resp_info['date'],
            'issuer_name': issuer_info['name'], 'issuer_job': issuer_info['job'], 'issuer_date': issuer_info['date'],
            'supervisor_name': super_info['name'], 'supervisor_job': super_info['job'], 'supervisor_date': super_info['date'],
            'work_place': permit.data.get('workPlace', permit.location.name if permit.location else ""),
            'work_name': permit.data.get('workName', ""),
            'work_content': permit.data.get('content', ""),
            'date_start': format_date(permit.valid_from),
            'date_end': format_date(permit.valid_to),
            'm5_1': permit.data.get('m5_1_stop', '____________________'),
            'm5_2': permit.data.get('m5_2_disconnect', '____________________'),
            'm5_3': permit.data.get('m5_3_install', '____________________'),
            'm5_4': permit.data.get('m5_4_analysis', '____________________'),
            'm5_5': permit.data.get('m5_5_fence', '____________________'),
            'm5_6': permit.data.get('m5_6_height', '____________________'),
            'm5_7': permit.data.get('m5_7_warn', '____________________'),
            'm5_8': permit.data.get('m5_8_railway', '____________________'),
            'm5_9': permit.data.get('m5_9_routes', '____________________'),
            'm5_10': permit.data.get('m5_10_additional', '____________________'),
            'brigade_list': brigade_list,
            'team_count': team_count,
            'extension_list': extension_list,
            'additional_coords': additional_coords_list,
            'has_additional_coords': len(additional_coords_list) > 0,
        }

        doc.render(context)

        # Исправляем высоту строк таблицы бригады: разрешаем строкам расти под размер подписи
        self._fix_brigade_row_heights(doc)

        # Оптимизация: основной наряд (13 пунктов) — на одну страницу
        self._compact_permit_to_one_page(doc)

        # ============================================================
        # ДОБАВЛЯЕМ ТАБЛИЦУ АНАЛИЗА РИСКОВ (перед чек-листами)
        # ============================================================
        self._append_risk_table_to_docx(doc, permit)

        # ============================================================
        # ДОБАВЛЯЕМ ЧЕК-ЛИСТЫ В ДОКУМЕНТ (после таблицы анализа рисков)
        # ============================================================
        self._append_checklists_to_docx(doc, permit)

        return doc

    def _convert_docx_to_pdf(self, docx_bytes):
        """
        Конвертирует DOCX (bytes) в PDF через LibreOffice.
        Возвращает bytes PDF или None при ошибке / отсутствии LibreOffice.
        """
        import subprocess
        import tempfile
        import os

        if not docx_bytes:
            return None
        tmpdir = tempfile.mkdtemp()
        try:
            docx_path = os.path.join(tmpdir, 'permit.docx')
            with open(docx_path, 'wb') as f:
                f.write(docx_bytes)
            # LibreOffice / OpenOffice: --headless --convert-to pdf --outdir <dir> <file>
            pdf_path = os.path.join(tmpdir, 'permit.pdf')
            for cmd_name in ('libreoffice', 'soffice'):
                try:
                    cmd = [
                        cmd_name, '--headless', '--convert-to', 'pdf',
                        '--outdir', tmpdir, docx_path
                    ]
                    result = subprocess.run(cmd, capture_output=True, timeout=60, cwd=tmpdir)
                    if result.returncode == 0 and os.path.isfile(pdf_path):
                        with open(pdf_path, 'rb') as f:
                            return f.read()
                except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
                    continue
            return None
        except Exception:
            return None
        finally:
            try:
                for f in os.listdir(tmpdir):
                    os.unlink(os.path.join(tmpdir, f))
                os.rmdir(tmpdir)
            except OSError:
                pass

    def _build_docx_response(self, permit, qr_url):
        """Собирает DOCX по наряду и возвращает HttpResponse (DOCX)."""
        from django.http import HttpResponse

        doc = self._get_rendered_doc(permit, qr_url)
        if doc is None:
            return None
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Permit_{permit.permit_id}.docx"'
        doc.save(response)
        return response

    # ============================================================
    # ЧЕКЛ-ЛИСТЫ В DOCX
    # ============================================================
    # Определение чек-листов (id -> название, список вопросов)
    CHECKLIST_DEFINITIONS = {
        'permit_check': {
            'title': 'ПРОВЕРКА ПО НАРЯДУ ДОПУСКУ',
            'required': True,
            'questions': {
                1: 'Пройдено ли персоналом обучение и имеются ли удостоверения?',
                2: 'Проведен ли инструктаж перед началом работы?',
                3: 'Обучены ли работники безопасным методам выполнения задания?',
                4: 'Подготовлены ли все необходимые разрешительные документы, включая НДПР?',
                5: 'Проведён ли инструктаж по действиям в случае пожара?',
                6: 'Имеет ли персонал соответствующую квалификацию?',
                7: 'Была ли проведена оценка рисков с определением всех опасностей?',
                8: 'Обеспечено ли правильное выявление источников опасности?',
                9: 'Проводилась ли новая оценка риска при изменении методов работ?',
                10: 'Правильно ли указано место производства работ в НД?',
                11: 'Все лица, выполняющие работу, понимают требования НД?',
                12: 'Правильно ли заполнен НДПР без корректирующих записей?',
                13: 'Учитывались ли все возможные опасности?',
                14: 'Получено ли разрешение перед блокировкой защитных ограждений?',
            },
        },
        'safety_bypass': {
            'title': 'ПРОВЕРКА ПО ОБХОДУ МЕР БЕЗОПАСНОСТИ',
            'required': True,
            'questions': {
                1: 'Используется ли оборудование для безопасности?',
                2: 'Ознакомлена ли бригада с планом аварийно-спасательных работ?',
                3: 'Были ли изменены или приспособлены СИЗ?',
                4: 'Обеспечено ли наличие средств пожаротушения?',
                5: 'Устранены ли все посторонние предметы?',
                6: 'Проверено ли наличие схемы распределения нагрузки?',
                7: 'Применяются ли защитные ограждения или барьеры?',
                8: 'Применяются ли средства защиты от движущихся частей?',
            },
        },
        'general_questions': {
            'title': 'ПРОВЕРКА ПО ОБЩИМ ВОПРОСАМ',
            'required': True,
            'questions': {
                1: 'Осуществляется ли регулярная проверка энергетических систем?',
                2: 'Зафиксированы ли все происшествия и аварийные ситуации?',
                3: 'Проводится ли медицинский осмотр работников?',
                4: 'Установлены ли зоны для безопасного хранения баллонов с газами?',
                5: 'Осуществляется ли контроль за состоянием скважин и трубопроводов?',
                6: 'Используются ли подъемные механизмы для тяжелых объектов?',
                7: 'Проверяется ли правильность крепления груза при подъеме?',
                8: 'Проводится ли проверка оборудования на наличие дефектов?',
            },
        },
        'confined_space': {
            'title': 'ПРОВЕРКА ПО РАБОТАМ В ЗАМКНУТОМ ПРОСТРАНСТВЕ',
            'required': False,
            'questions': {
                1: 'Проведена ли оценка замкнутого объема?',
                2: 'Соответствует ли содержание вредных веществ нормам?',
                3: 'Соответствует ли частота отбора проб воздушной среды?',
                4: 'Обеспечена ли вентиляция в замкнутом объеме?',
                5: 'Назначен ли ответственный наблюдатель по ТБ?',
                6: 'Обеспечены ли работники спасательным оборудованием?',
                7: 'Проверено ли электрооборудование в замкнутом объеме?',
                8: 'Котлованы глубиной более 1,2 м имеют оборудованные места входа?',
            },
        },
        'lifting_works': {
            'title': 'ПРОВЕРКА ПО ГРУЗОПОДЪЕМНЫМ РАБОТАМ',
            'required': False,
            'questions': {
                1: 'Проверены ли средства защиты от падения?',
                2: 'Ознакомлена ли бригада с планом аварийно-спасательных работ?',
                3: 'Отрегулированы ли средства защиты от падения?',
                4: 'Ознакомлены ли члены бригады с правилами использования средств защиты?',
                5: 'Используют ли работники 100% крепление ИСС?',
                6: 'Проведена ли оценка риска падения предметов?',
                7: 'Проинформированы ли члены бригады о запрете работы без напарника?',
                8: 'Имеются ли средства защиты (ограждения, сетки)?',
                9: 'Правильно ли возведены строительные леса?',
                10: 'Безопасно ли установлены переносные лестницы?',
                11: 'Утвержден ли письменный план грузоподъемных работ?',
                12: 'Подтверждена ли квалификация крановщика и стропальщика?',
                13: 'Проведена ли проверка веса груза на соответствие ГПМ?',
                14: 'Соответствуют ли ГПМ и такелажное оборудование требованиям?',
                15: 'Предохранительные устройства на ГПМ в рабочем состоянии?',
                16: 'Согласован ли план взаимодействия всеми членами бригады?',
                17: 'Разработан ли план по предупреждению опасных факторов?',
                18: 'ГПМ находится в устойчивом положении?',
                19: 'Имелись ли средства защиты для предотвращения входа?',
                20: 'Подъемное устройство расположено без препятствий?',
                21: 'Понимает ли оператор как осуществить аварийный останов?',
            },
        },
        'fire_works': {
            'title': 'ПРОВЕРКА ИНСПЕКТОРА ПО ОГНЕВЫМ РАБОТАМ',
            'required': False,
            'questions': {
                1: 'Соблюдены ли требования по изоляции источников энергии?',
                2: 'Проводился ли отбор проб воздушной среды?',
                3: 'Установлена ли частота отбора проб?',
                4: 'Проведена ли проверка герметичности?',
                5: 'Источники возгорания перемещены на безопасное расстояние?',
                6: 'Имеются ли средства пожаротушения?',
                7: 'Участок огражден от разлета тепла и искр?',
            },
        },
        'energy_isolation': {
            'title': 'ПРОВЕРКА ПО ИЗОЛЯЦИИ ОПАСНЫХ ЭНЕРГИЙ',
            'required': False,
            'questions': {
                1: 'Изолирующие устройства установлены по схеме?',
                2: 'Изолирующие устройства соответствуют требованиям?',
                3: 'На точках изоляции вывешены замки и ярлыки?',
                4: 'Проведена ли проверка остаточной энергии?',
                5: 'Определено ли нахождение работников "под ударом"?',
                6: 'Дренажные и воздушные клапаны в открытом положении?',
                7: 'Проверено ли соответствие точки разгерметизации?',
                8: 'Проведена ли снятие устройств изоляции по плану?',
            },
        },
        'danger_zone': {
            'title': 'ПРОВЕРКА ПО ОПАСНОЙ ЗОНЕ',
            'required': False,
            'questions': {
                1: 'Имелись ли ограждения для предотвращения доступа?',
                2: 'Установлены ли предупреждающие знаки?',
                3: 'Опасные зоны определены и доступ ограничен?',
                4: 'Исполнители находились в безопасном месте?',
                5: 'Определены и ограждены безопасные рабочие участки?',
                6: 'Токоведущие части ограждены или изолированы?',
                7: 'Осуществляется ли обмен информацией?',
            },
        },
        'vehicle_driving': {
            'title': 'ПРОВЕРКА ПО ВОЖДЕНИЮ ТРАНСПОРТНОГО СРЕДСТВА',
            'required': False,
            'questions': {
                1: 'Водитель имеет действующее водительское удостоверение?',
                2: 'ТС оборудовано средствами безопасности?',
                3: 'Имеется штамп о медицинском освидетельствовании?',
                4: 'В исправном ли состоянии фары и стоп-сигналы?',
                5: 'Закреплен ли груз?',
                6: 'Имеется разрешение на перевозку опасных грузов?',
                7: 'Водитель выполняет требования безопасности?',
                8: 'Имеются все действительные документы?',
                9: 'Чисто ли ТС и видны ли средства идентификации?',
                10: 'Имеется ли план движения по маршруту?',
            },
        },
    }

    @staticmethod
    def _fix_brigade_row_heights(doc):
        """Разрешает строкам таблицы бригады расширяться под размер подписи (atLeast вместо exact)."""
        from lxml import etree
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        for table in doc.docx.tables:
            for row in table.rows:
                # Ищем строки с изображением подписи (после рендеринга содержат <a:blip> внутри)
                tr_xml = etree.tostring(row._tr)
                has_image = b'<a:blip' in tr_xml or b'graphicData' in tr_xml
                if has_image:
                    trPr = row._tr.find(f'{{{ns}}}trPr')
                    if trPr is None:
                        trPr = etree.SubElement(row._tr, f'{{{ns}}}trPr')
                        row._tr.insert(0, trPr)
                    trHeight = trPr.find(f'{{{ns}}}trHeight')
                    if trHeight is not None:
                        trHeight.set(f'{{{ns}}}hRule', 'atLeast')
                    else:
                        trHeight = etree.SubElement(trPr, f'{{{ns}}}trHeight')
                        trHeight.set(f'{{{ns}}}val', '500')
                        trHeight.set(f'{{{ns}}}hRule', 'atLeast')

    def _compact_permit_to_one_page(self, doc):
        """
        Оптимизация: уменьшает поля, отступы и межстрочные интервалы основного наряда,
        чтобы все 13 пунктов поместились на одну страницу.
        """
        from docx.shared import Pt, Cm

        document = doc.docx
        if not document.sections:
            return

        # Уменьшаем поля страницы (по умолчанию ~1.25 см — делаем ~1 см)
        section = document.sections[0]
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

        def compact_paragraph(paragraph):
            pf = paragraph.paragraph_format
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            pf.line_spacing = Pt(10)  # межстрочный
            for run in paragraph.runs:
                try:
                    if run.font.size is not None and run.font.size.pt > 10:
                        run.font.size = Pt(9)
                except (AttributeError, TypeError):
                    pass

        from lxml import etree
        wml_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

        def _has_image(element):
            raw = etree.tostring(element)
            return b'<a:blip' in raw or b'graphicData' in raw

        def _fix_image_paragraph(paragraph):
            """Убирает lineRule=exact, ставит atLeast — чтобы строка расширилась под картинку."""
            pf = paragraph.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pPr = paragraph._p.find(f'{wml_ns}pPr')
            if pPr is not None:
                spacing = pPr.find(f'{wml_ns}spacing')
                if spacing is not None:
                    lr = spacing.get(f'{wml_ns}lineRule')
                    if lr == 'exact':
                        spacing.set(f'{wml_ns}lineRule', 'atLeast')

        # Компактные стили для всех параграфов в теле документа
        for paragraph in document.paragraphs:
            if _has_image(paragraph._p):
                _fix_image_paragraph(paragraph)
            else:
                compact_paragraph(paragraph)

        # Параграфы внутри таблиц (пропускаем ячейки с изображениями подписей)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if _has_image(cell._tc):
                        for paragraph in cell.paragraphs:
                            _fix_image_paragraph(paragraph)
                    else:
                        for paragraph in cell.paragraphs:
                            compact_paragraph(paragraph)

    def _append_risk_table_to_docx(self, doc, permit):
        """
        Добавляет "Таблицу анализа рисков" из data.riskTable перед чек-листами.
        """
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        risk_rows = permit.data.get('riskTable', []) if permit.data else []
        if not risk_rows or not isinstance(risk_rows, list):
            return

        # Оставляем только строки с реальными данными
        rows = []
        for r in risk_rows:
            if not isinstance(r, dict):
                continue
            if any((r.get('step'), r.get('hazards'), r.get('measures'), r.get('isControlled'))):
                rows.append(r)

        if not rows:
            return

        document = doc.docx
        document.add_page_break()

        title_p = document.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run('ТАБЛИЦА АНАЛИЗА РИСКОВ')
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_run.font.name = 'Times New Roman'

        spacer = document.add_paragraph()
        spacer.space_after = Pt(6)

        tbl = document.add_table(rows=len(rows) + 1, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = 'Table Grid'

        header = ['№', 'Шаг работы', 'Опасности', 'Меры контроля', 'Контроль (Да/Нет)']
        widths = [Cm(1), Cm(4.2), Cm(5.2), Cm(5.2), Cm(2.0)]

        for i, text in enumerate(header):
            cell = tbl.rows[0].cells[i]
            cell.text = text
            cell.width = widths[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'

        for idx, row in enumerate(rows, start=1):
            values = [
                str(idx),
                str(row.get('step') or ''),
                str(row.get('hazards') or ''),
                str(row.get('measures') or ''),
                str(row.get('isControlled') or ''),
            ]
            for col, val in enumerate(values):
                cell = tbl.rows[idx].cells[col]
                cell.text = val
                cell.width = widths[col]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT
                run = p.runs[0] if p.runs else p.add_run(val)
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'
                if col == 4:
                    up = val.strip().upper()
                    if up in ('ДА', 'YES', 'Y', '+'):
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                    elif up in ('НЕТ', 'NO', 'N', '-'):
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    def _append_checklists_to_docx(self, doc, permit):
        """
        Добавляет заполненные чек-листы как новые страницы в DOCX после рендеринга шаблона.
        Добавляются только те чек-листы, которые были заполнены пользователем.
        """
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn

        checklist_data = permit.data.get('checklist', {})
        if not checklist_data:
            return

        ANSWER_MAP = {'YES': 'ДА', 'NO': 'НЕТ', 'NA': 'Н/П'}

        # Порядок: сначала обязательные, потом остальные
        ordered_ids = [
            'permit_check', 'safety_bypass', 'general_questions',
            'confined_space', 'lifting_works', 'fire_works',
            'energy_isolation', 'danger_zone', 'vehicle_driving',
        ]

        document = doc.docx  # Доступ к python-docx Document из docxtpl

        # Все заполненные чек-листы на одной странице: один разрыв перед секцией, между таблицами — без разрыва
        first_checklist = True

        for table_id in ordered_ids:
            table_answers = checklist_data.get(table_id)
            if not table_answers:
                continue

            # Пропускаем, если ни один вопрос не отвечен
            has_any_answer = any(
                a.get('answer', '') != ''
                for a in table_answers.values() if isinstance(a, dict)
            )
            if not has_any_answer:
                continue

            table_def = self.CHECKLIST_DEFINITIONS.get(table_id)
            if not table_def:
                continue

            # --- Одна новая страница только перед первым чек-листом ---
            if first_checklist:
                document.add_page_break()
                first_checklist = False
            else:
                # Между чек-листами — небольшой отступ, без разрыва страницы
                spacer_between = document.add_paragraph()
                spacer_between.space_before = Pt(12)
                spacer_between.space_after = Pt(6)

            # --- ЗАГОЛОВОК ---
            title_p = document.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run(table_def['title'])
            title_run.bold = True
            title_run.font.size = Pt(12)
            title_run.font.name = 'Times New Roman'

            if table_def.get('required'):
                req_p = document.add_paragraph()
                req_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                req_run = req_p.add_run('(Обязательный)')
                req_run.bold = True
                req_run.font.size = Pt(9)
                req_run.font.name = 'Times New Roman'
                req_run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

            # Пустая строка
            spacer = document.add_paragraph()
            spacer.space_after = Pt(4)

            # --- ТАБЛИЦА ---
            questions = table_def['questions']
            num_rows = len(questions) + 1  # +1 для заголовка
            tbl = document.add_table(rows=num_rows, cols=5)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.style = 'Table Grid'

            # Ширины столбцов
            widths = [Cm(1), Cm(10), Cm(1.5), Cm(1.5), Cm(4)]

            # --- ЗАГОЛОВОК ТАБЛИЦЫ ---
            header_cells = tbl.rows[0].cells
            header_texts = ['№', 'ДЕЙСТВИЯ', 'ДА/НЕТ/Н/П', 'Ответ', 'Комментарии']
            for i, (cell, text) in enumerate(zip(header_cells, header_texts)):
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'
                # Задаем ширину
                cell.width = widths[i]
                # Цвет фона заголовка
                shading = cell._element.get_or_add_tcPr()
                shd_elem = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'D9E2F3',
                })
                shading.append(shd_elem)

            # --- СТРОКИ ДАННЫХ ---
            for row_idx, (q_num, q_text) in enumerate(questions.items(), start=1):
                row_cells = tbl.rows[row_idx].cells

                # Получаем ответ
                q_answer_data = table_answers.get(str(q_num), {})
                if isinstance(q_answer_data, dict):
                    answer_code = q_answer_data.get('answer', '')
                    comment = q_answer_data.get('comment', '')
                else:
                    answer_code = ''
                    comment = ''

                answer_text = ANSWER_MAP.get(answer_code, '—')

                row_data = [str(q_num), q_text, '', answer_text, comment]

                for i, (cell, text) in enumerate(zip(row_cells, row_data)):
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run(text)
                    run.font.size = Pt(8)
                    run.font.name = 'Times New Roman'
                    cell.width = widths[i]

                    if i == 0:  # Номер — по центру
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif i == 3:  # Ответ — по центру, с цветом
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.bold = True
                        if answer_code == 'YES':
                            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                        elif answer_code == 'NO':
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

            # Информация о наряде внизу таблицы
            footer_p = document.add_paragraph()
            footer_p.space_before = Pt(8)
            footer_run = footer_p.add_run(f'Наряд-допуск №{permit.permit_id}')
            footer_run.font.size = Pt(8)
            footer_run.font.name = 'Times New Roman'
            footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    @action(detail=True, methods=['get'], url_path='download_docx')
    def download_docx(self, request, pk=None):
        """Скачивание наряда в формате PDF (конвертация из DOCX). При отсутствии LibreOffice отдаётся DOCX."""
        from django.http import HttpResponse
        from django.conf import settings
        from io import BytesIO

        permit = self.get_object()
        base = getattr(settings, 'HSE_BASE_URL', 'https://hse.kbm.kz')
        qr_url = f"{base}/api/v1/verify/{permit.verify_token}/"
        try:
            doc = self._get_rendered_doc(permit, qr_url)
            if doc is None:
                return Response({"error": "Шаблон не найден"}, status=500)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_bytes = buffer.getvalue()

            pdf_bytes = self._convert_docx_to_pdf(docx_bytes)
            if pdf_bytes:
                response = HttpResponse(pdf_bytes, content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="Permit_{permit.permit_id}.pdf"'
                return response

            # Fallback: отдаём DOCX, если LibreOffice недоступен
            response = HttpResponse(docx_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="Permit_{permit.permit_id}.docx"'
            return response
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"error": f"Ошибка генерации документа: {str(e)}"}, status=500)

    @action(detail=True, methods=['get'], url_path='verify_docx', permission_classes=[AllowAny])
    def verify_docx(self, request, pk=None):
        """Редирект старых QR-кодов на новый URL с токеном."""
        from django.http import HttpResponseRedirect
        permit = get_object_or_404(WorkPermit, pk=pk)
        return HttpResponseRedirect(f'/api/v1/verify/{permit.verify_token}/')

    def destroy(self, request, *args, **kwargs):
        user = request.user
        if user.is_auditor:
            return Response({'error': 'Аудитор не может удалять наряды.'}, status=403)
        permit = self.get_object()
        if permit.status != 'DRAFT' or permit.initiator != user:
            if not user.is_admin:
                return Response({'error': 'Удалять можно только свои черновики.'}, status=403)
        return super().destroy(request, *args, **kwargs)

    def perform_update(self, serializer):
        """
        Логика сохранения:
        - Инициатор (Выдающий наряд) может редактировать черновик без ограничений.
        - Согласующие (Ответственный/Допускающий/Производитель) могут редактировать
          во время согласования, но подписанты защищены на сервере —
          роли (issuer, responsible, admitting, producer, supervisor) не меняются.
        """
        permit = self.get_object()
        user = self.request.user

        if user.is_auditor:
            raise PermissionDenied("Аудитор не может редактировать наряды.")

        # 1. Инициатор (Выдающий наряд) редактирует черновик или отклонённый наряд — без ограничений
        if permit.initiator == user and permit.status in ('DRAFT', 'REJECTED'):
            serializer.save()
            return

        # 2. Согласующий редактирует во время согласования
        try:
            current_step = permit.approval_steps.get(status='PENDING', approver=user)
            allowed_roles = ['RESPONSIBLE', 'ADMITTING', 'WORK_PRODUCER']

            if permit.status == 'PENDING_APPROVAL' and current_step.role in allowed_roles:
                # Защита: принудительно сохраняем ОРИГИНАЛЬНЫЕ роли подписантов,
                # чтобы согласующий не мог их изменить
                old_data = permit.data or {}
                new_data = serializer.validated_data.get('data', {})

                role_keys = ['issuer', 'responsible', 'admitting', 'producer', 'supervisor']
                for key in role_keys:
                    if key in old_data:
                        new_data[key] = old_data[key]

                serializer.validated_data['data'] = new_data
                serializer.save()
                return

        except ApprovalStep.DoesNotExist:
            pass

        # Если ни одно условие не совпало — запрещаем
        raise PermissionDenied("У вас нет прав на редактирование наряда в текущем статусе.")

    def perform_create(self, serializer):
        user = self.request.user

        # 👇 ПРОВЕРКА РОЛИ: создавать наряд могут только Выдающий наряд (ISSUER), Допускающий (ADMITTING) или Админ
        if user.role not in ('ISSUER', 'ADMITTING') and not user.is_admin:
            raise PermissionDenied("Создавать наряды могут только пользователи с ролями «Выдающий наряд» или «Допускающий».")

        # Если проверка пройдена, сохраняем (ваш старый код инициатора)
        serializer.save(initiator=user)

    @action(detail=True, methods=['post'], url_path='upload_safety_document')
    def upload_safety_document(self, request, pk=None):
        """
        Прикрепить документ к мерам безопасности (PDF или JPG).
        Доступно при создании/редактировании наряда (DRAFT, REJECTED) или во время согласования.
        """
        permit = self.get_object()
        user = request.user

        # Право загрузки: инициатор черновика/отклонённого или согласующий с правом редактирования
        can_edit = (
            (permit.status in ('DRAFT', 'REJECTED') and permit.initiator == user) or
            (permit.status == 'PENDING_APPROVAL' and permit.approval_steps.filter(
                approver=user, status='PENDING', role__in=['RESPONSIBLE', 'ADMITTING', 'WORK_PRODUCER']
            ).exists()) or
            user.is_admin
        )
        if not can_edit:
            return Response({'error': 'Нет прав на прикрепление документа.'}, status=403)

        doc_file = request.FILES.get('safety_document')
        if not doc_file:
            return Response({'error': 'Не выбран файл. Прикрепите PDF или JPG.'}, status=400)

        allowed_extensions = ('.pdf', '.jpg', '.jpeg')
        file_ext = os.path.splitext(doc_file.name)[1].lower()
        if file_ext not in allowed_extensions:
            return Response({
                'error': f'Недопустимый формат ({file_ext}). Разрешены только: PDF, JPG.'
            }, status=400)

        max_size_mb = 10
        if doc_file.size > max_size_mb * 1024 * 1024:
            file_size_mb = round(doc_file.size / (1024 * 1024), 1)
            return Response({
                'error': f'Файл слишком большой: {file_size_mb} МБ. Максимум: {max_size_mb} МБ.'
            }, status=400)

        permit.safety_document = doc_file
        permit.save(update_fields=['safety_document'])

        return Response({
            'status': 'Документ прикреплён.',
            'safety_document_url': permit.safety_document.url if permit.safety_document else None
        })

    @action(detail=True, methods=['post'], url_path='upload_loto_photo')
    def upload_loto_photo(self, request, pk=None):
        permit = self.get_object()
        user = request.user

        can_edit = (
            (permit.status in ('DRAFT', 'REJECTED') and permit.initiator == user) or
            (permit.status == 'PENDING_APPROVAL' and permit.approval_steps.filter(
                approver=user, status='PENDING', role__in=['RESPONSIBLE', 'ADMITTING', 'WORK_PRODUCER']
            ).exists()) or
            user.is_admin
        )
        if not can_edit:
            return Response({'error': 'Нет прав на загрузку файла.'}, status=403)

        photo_file = request.FILES.get('loto_photo')
        if not photo_file:
            return Response({'error': 'Файл не выбран.'}, status=400)

        allowed_extensions = ('.pdf', '.jpg', '.jpeg', '.png')
        file_ext = os.path.splitext(photo_file.name)[1].lower()
        if file_ext not in allowed_extensions:
            return Response({
                'error': f'Недопустимый формат ({file_ext}). Разрешены: PDF, JPG, PNG.'
            }, status=400)

        max_size_mb = 10
        if photo_file.size > max_size_mb * 1024 * 1024:
            return Response({
                'error': f'Файл слишком большой. Максимум: {max_size_mb} МБ.'
            }, status=400)

        permit.loto_photo = photo_file
        permit.save(update_fields=['loto_photo'])

        return Response({
            'status': 'Фото загружено.',
            'loto_photo_url': permit.loto_photo.url
        })

    @action(detail=True, methods=['delete'], url_path='delete_loto_photo')
    def delete_loto_photo(self, request, pk=None):
        permit = self.get_object()
        user = request.user

        can_edit = (
            (permit.status in ('DRAFT', 'REJECTED') and permit.initiator == user) or
            (permit.status == 'PENDING_APPROVAL' and permit.approval_steps.filter(
                approver=user, status='PENDING', role__in=['RESPONSIBLE', 'ADMITTING', 'WORK_PRODUCER']
            ).exists()) or
            user.is_admin
        )
        if not can_edit:
            return Response({'error': 'Нет прав на удаление файла.'}, status=403)

        if permit.loto_photo:
            permit.loto_photo.delete(save=False)
            permit.loto_photo = None
            permit.save(update_fields=['loto_photo'])

        return Response({'status': 'Фото удалено.'})

    @action(detail=True, methods=['post'], url_path='producer_close')
    def producer_close(self, request, pk=None):
        """
        Шаг 1 закрытия: Производитель работ подтверждает завершение.
        Если исполнитель внешний (без ЭЦП), подтвердить могут Выдающий или Допускающий.
        """
        permit = self.get_object()
        user = request.user

        if permit.status != 'APPROVED':
            return Response({'error': 'Закрыть можно только согласованный наряд.'}, status=400)

        if permit.producer_closed:
            return Response({'error': 'Производитель работ уже подтвердил закрытие.'}, status=400)

        producer_step = permit.approval_steps.filter(role='WORK_PRODUCER').first()

        # Обычный производитель — есть учётная запись
        is_producer = (
            producer_step and producer_step.approver_id and
            str(producer_step.approver_id) == str(user.id)
        )

        # Внешний производитель (без ЭЦП) — действует Выдающий или Допускающий
        external_producer = (
            producer_step and not producer_step.approver_id and
            isinstance((permit.data or {}).get('producer'), dict) and
            permit.data['producer'].get('external')
        )
        is_issuer_or_admitting = permit.approval_steps.filter(
            role__in=['ISSUER', 'ADMITTING'], approver=user, status='APPROVED'
        ).exists()

        can_close = is_producer or (external_producer and is_issuer_or_admitting) or user.is_admin

        if not can_close:
            return Response(
                {'error': 'Только Производитель работ может подтвердить завершение работ.'},
                status=403
            )

        # Проверка: все члены бригады должны подписать наряд перед закрытием
        permit_data = permit.data or {}
        team = permit_data.get('teamMembers') or []
        if team:
            sigs = permit_data.get('brigade_signatures')
            if isinstance(sigs, dict):
                sigs = [sigs.get(str(i)) for i in range(len(team))]
            elif not isinstance(sigs, list):
                sigs = []
            unsigned = [
                {'index': i, 'name': (m.get('name') or '—')}
                for i, m in enumerate(team)
                if i >= len(sigs) or not sigs[i]
            ]
            if unsigned:
                names = '; '.join(f"№{u['index'] + 1} {u['name']}" for u in unsigned)
                return Response(
                    {
                        'error': (
                            'Вы не можете закрыть наряд, пока не подпишут «Состав бригады». '
                            f'Не подписали: {names}.'
                        ),
                        'unsigned_members': unsigned,
                    },
                    status=400,
                )

        # Сохраняем графическую подпись производителя при закрытии (если внешний)
        signature_file = request.FILES.get('signature')
        if external_producer and signature_file:
            rel_dir = os.path.join('brigade_signatures', str(permit.pk))
            dest_dir = os.path.join(settings.MEDIA_ROOT, rel_dir)
            os.makedirs(dest_dir, exist_ok=True)
            fname = 'producer_close.png'
            rel_path = os.path.join(rel_dir, fname).replace('\\', '/')
            full_path = os.path.join(settings.MEDIA_ROOT, rel_dir, fname)
            with open(full_path, 'wb') as f:
                for chunk in signature_file.chunks():
                    f.write(chunk)
            data = dict(permit.data) if permit.data else {}
            data['producer_close_signature'] = rel_path
            permit.data = data
            permit.save(update_fields=['data'])

        permit.producer_closed = True
        permit.save(update_fields=['producer_closed'])

        # Уведомляем Допускающего
        admitting_step = permit.approval_steps.filter(role='ADMITTING', status='APPROVED').first()
        if admitting_step and admitting_step.approver:
            Notification.objects.create(
                user=admitting_step.approver,
                permit_id=permit.id,
                title="Закройте наряд",
                message=(
                    f"Производитель работ подписал и подтвердил завершение работ по наряду №{permit.permit_id}. "
                    f"Теперь вам необходимо закрыть наряд (кнопка «Закрыть наряд как Допускающий»)."
                )
            )

        return Response({'ok': True, 'message': 'Производитель работ подтвердил завершение работ.'})

    @action(detail=True, methods=['post'], url_path='close')
    def close_permit(self, request, pk=None):
        """
        Шаг 2 закрытия: Допускающий окончательно закрывает наряд.
        Требует, чтобы Производитель работ уже подтвердил закрытие (producer_closed=True).
        """
        permit = self.get_object()
        user = request.user

        # 1. Проверка статуса
        if permit.status != 'APPROVED':
            return Response({'error': 'Закрыть можно только согласованный наряд.'}, status=400)

        # 2. Производитель должен подтвердить первым
        if not permit.producer_closed:
            return Response(
                {'error': 'Сначала Производитель работ должен подтвердить завершение работ.'},
                status=400
            )

        # 3. Проверка прав (только Допускающий)
        is_admitting = permit.approval_steps.filter(role='ADMITTING', approver=user).exists()
        if not is_admitting:
            admitting_data = (permit.data or {}).get('admitting') or {}
            if str(admitting_data.get('id')) == str(user.id):
                is_admitting = True

        if not is_admitting and not user.is_admin:
            return Response({'error': 'Только Допускающий имеет право закрыть наряд.'}, status=403)

        # 4. Закрываем наряд
        permit.valid_to = timezone.now()
        permit.close_work()
        permit.save()

        return Response({'ok': True, 'status': 'Наряд успешно закрыт'})






class DepartmentViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Справочник департаментов (только чтение для API)
    """
    queryset = Department.objects.all()
    serializer_class = DepartamentSerializer
    permission_classes = [IsAuthenticated]

    filter_backends = [filters.SearchFilter]
    search_fields = ['name', 'name_kk']


class DangerousWorkTypeViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Справочник опасных работ
    """
    queryset = DangerousWorkType.objects.all()
    serializer_class = DangerousWorkTypeSerializer
    permission_classes = [IsAuthenticated]

    filter_backends = [filters.SearchFilter]
    search_fields = ['name', 'name_kk']


# 👇 НОВЫЙ VIEWSET ДЛЯ УВЕДОМЛЕНИЙ (Колокольчик)
class NotificationViewSet(viewsets.ReadOnlyModelViewSet):
    permission_classes = [IsAuthenticated]
    serializer_class = NotificationSerializer

    def get_queryset(self):
        # Показываем только уведомления текущего юзера
        return Notification.objects.filter(user=self.request.user).order_by('-created_at')

    @action(detail=True, methods=['post'])
    def mark_read(self, request, pk=None):
        notif = self.get_object()
        notif.is_read = True
        notif.save()
        return Response({'status': 'ok'})


def verify_permit_public(request, token):
    """Публичная верификация наряда по QR-токену. Без авторизации, только чтение."""
    from django.http import HttpResponse
    from django.shortcuts import get_object_or_404
    from zoneinfo import ZoneInfo

    permit = get_object_or_404(WorkPermit, verify_token=token)
    kz_tz = ZoneInfo('Asia/Almaty')

    def fmt(dt):
        if not dt:
            return '—'
        return dt.astimezone(kz_tz).strftime('%d.%m.%Y %H:%M')

    status_map = {
        'DRAFT': ('Черновик', '#6b7280'),
        'PENDING_APPROVAL': ('На согласовании', '#f59e0b'),
        'APPROVED': ('Согласован', '#10b981'),
        'REJECTED': ('Отклонён', '#ef4444'),
        'CLOSED': ('Закрыт', '#6366f1'),
    }
    status_label, status_color = status_map.get(permit.status, (permit.status, '#6b7280'))

    steps_html = ''
    for step in permit.approval_steps.select_related('approver').order_by('step_order'):
        approver_name = step.approver.get_full_name() if step.approver else '—'
        role_label = step.get_role_display() if hasattr(step, 'get_role_display') else step.role
        step_status_map = {
            'PENDING': ('Ожидает', '#f59e0b', '&#9203;'),
            'APPROVED': ('Подписан', '#10b981', '&#9989;'),
            'REJECTED': ('Отклонён', '#ef4444', '&#10060;'),
        }
        s_label, s_color, s_icon = step_status_map.get(step.status, (step.status, '#6b7280', ''))
        signed_str = fmt(step.signed_at) if step.signed_at else ''
        steps_html += f'''
        <tr>
            <td>{role_label}</td>
            <td>{approver_name}</td>
            <td><span class="status-badge" style="background:{s_color}">{s_icon} {s_label}</span></td>
            <td class="date-col">{signed_str}</td>
        </tr>'''

    location_name = permit.location.name if permit.location else '—'
    work_name = (permit.data or {}).get('workName', '—')
    work_place = (permit.data or {}).get('workPlace', location_name)
    initiator_name = permit.initiator.get_full_name() if permit.initiator else '—'

    html = f'''<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Верификация наряда {permit.permit_id}</title>
    <style>
        * {{ margin:0; padding:0; box-sizing:border-box; }}
        body {{ font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif; background:#f1f5f9; color:#1e293b; }}
        .container {{ max-width:700px; margin:30px auto; padding:0 16px; }}
        .card {{ background:#fff; border-radius:16px; box-shadow:0 1px 3px rgba(0,0,0,.1); overflow:hidden; }}
        .header {{ background:linear-gradient(135deg,#1e3a5f,#2563eb); color:#fff; padding:28px 24px; text-align:center; }}
        .header h1 {{ font-size:20px; font-weight:700; margin-bottom:4px; }}
        .header p {{ font-size:14px; opacity:.85; }}
        .badge {{ display:inline-block; padding:6px 18px; border-radius:20px; font-weight:600; font-size:15px; margin-top:12px; color:#fff; background:{status_color}; }}
        .body {{ padding:24px; }}
        .section {{ margin-bottom:20px; }}
        .section-title {{ font-size:13px; font-weight:700; text-transform:uppercase; letter-spacing:.5px; color:#64748b; margin-bottom:10px; }}
        .grid {{ display:grid; grid-template-columns:1fr 1fr; gap:12px; }}
        @media(max-width:500px) {{ .grid {{ grid-template-columns:1fr; }} }}
        .field {{ background:#f8fafc; border-radius:10px; padding:12px 14px; }}
        .field .label {{ font-size:12px; color:#94a3b8; margin-bottom:3px; }}
        .field .value {{ font-size:15px; font-weight:500; }}
        .table-wrapper {{ overflow-x:auto; -webkit-overflow-scrolling:touch; }}
        table {{ width:100%; border-collapse:collapse; font-size:14px; min-width:500px; }}
        th {{ background:#f8fafc; padding:10px 12px; text-align:left; font-size:12px; text-transform:uppercase; letter-spacing:.5px; color:#64748b; border-bottom:2px solid #e2e8f0; white-space:nowrap; }}
        td {{ padding:10px 12px; border-bottom:1px solid #e2e8f0; }}
        .date-col {{ color:#64748b; font-size:13px; white-space:nowrap; }}
        .status-badge {{ color:#fff; padding:3px 10px; border-radius:12px; font-size:13px; white-space:nowrap; }}
        .footer {{ text-align:center; padding:16px; color:#94a3b8; font-size:12px; border-top:1px solid #f1f5f9; }}
        .verified {{ display:flex; align-items:center; justify-content:center; gap:8px; padding:14px; background:#f0fdf4; border-radius:10px; margin-bottom:20px; color:#16a34a; font-weight:600; font-size:15px; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="card">
            <div class="header">
                <h1>Наряд-допуск {permit.permit_id}</h1>
                <p>Система HSE — КБМ</p>
                <div class="badge">{status_label}</div>
            </div>
            <div class="body">
                <div class="verified">&#9989; Документ зарегистрирован в системе HSE KBM</div>

                <div class="section">
                    <div class="section-title">Основная информация</div>
                    <div class="grid">
                        <div class="field"><div class="label">Номер наряда</div><div class="value">{permit.permit_id}</div></div>
                        <div class="field"><div class="label">Статус</div><div class="value" style="color:{status_color}">{status_label}</div></div>
                        <div class="field"><div class="label">Инициатор</div><div class="value">{initiator_name}</div></div>
                        <div class="field"><div class="label">Подразделение</div><div class="value">{location_name}</div></div>
                        <div class="field"><div class="label">Место работ</div><div class="value">{work_place}</div></div>
                        <div class="field"><div class="label">Наименование работ</div><div class="value">{work_name}</div></div>
                        <div class="field"><div class="label">Начало работ</div><div class="value">{fmt(permit.valid_from)}</div></div>
                        <div class="field"><div class="label">Окончание работ</div><div class="value">{fmt(permit.valid_to)}</div></div>
                        <div class="field"><div class="label">Дата создания</div><div class="value">{fmt(permit.created_at)}</div></div>
                    </div>
                </div>

                <div class="section">
                    <div class="section-title">Согласование</div>
                    <div class="table-wrapper">
                        <table>
                            <thead><tr><th>Роль</th><th>ФИО</th><th>Статус</th><th>Дата</th></tr></thead>
                            <tbody>{steps_html if steps_html else '<tr><td colspan="4" style="padding:14px;text-align:center;color:#94a3b8;">Нет данных</td></tr>'}</tbody>
                        </table>
                    </div>
                </div>
            </div>
            <div class="footer">Система HSE KBM &bull; hse.kbm.kz</div>
        </div>
    </div>
</body>
</html>'''

    return HttpResponse(html, content_type='text/html; charset=utf-8')


class AIAssistantView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        user_question = request.data.get('question', '')

        if not user_question:
            return Response({"error": "Вопрос не может быть пустым"}, status=400)

        # 👇 Вставьте сюда ваш ключ от DeepSeek
        deepseek_api_key = settings.DEEPSEEK_API_KEY

        if not deepseek_api_key:
            return Response({"answer": "Ошибка: API ключ DeepSeek не настроен."}, status=200)

        # 👇 НАСТРОЙКА КЛИЕНТА ПОД DEEPSEEK
        client = OpenAI(
            api_key=deepseek_api_key,
            base_url="https://api.deepseek.com"  # Это перенаправляет запросы на сервера DeepSeek
        )

        # Инструкция, чтобы он не болтал лишнего
        system_prompt = """
        Ты — умный ассистент по технике безопасности (HSE) в нефтяной компании АО «Каражанбасмунай».
        Твоя задача — помогать персоналу грамотно заполнять наряды-допуски и оценивать риски.

        Правила:
        1. Отвечай ТОЛЬКО на вопросы, касающиеся безопасности, нарядов, рисков, СИЗ и промышленных работ.
        2. Если вопрос не по теме (погода, стихи, код), вежливо откажись: "Я консультирую только по вопросам производственной безопасности."
        3. Будь краток и давай четкие инструкции. Используй списки.
        4. Если спрашивают про меры безопасности, перечисли конкретные действия (оградить, отключить, проверить газоанализатором).
        """

        try:
            response = client.chat.completions.create(
                model="deepseek-chat",  # 👈 Используем их модель (V3)
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_question}
                ],
                temperature=0.3,  # Низкая температура для точности
                max_tokens=500
            )

            ai_answer = response.choices[0].message.content
            return Response({"answer": ai_answer})

        except Exception as e:
            print(f"DeepSeek Error: {e}")
            return Response({"answer": "Извините, сервис ИИ временно недоступен. Попробуйте позже."}, status=200)