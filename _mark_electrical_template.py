"""
Размечает «Бланк наряда на электроустановку (каз-рус).docx» под docxtpl —
заменяет подчёркивания «____» на jinja-теги {{ tag }} и {% for ... %}.
Результат сохраняется в templates/docx/electrical_permit_rus_kk.docx
"""
import os
import re
import shutil
from docx import Document
from copy import deepcopy

SRC = '/home/sadmin/web/dev/form_electric/Бланк наряда на электроустановку (каз-рус).docx'
DST = '/home/sadmin/web/dev/templates/docx/electrical_permit_rus_kk.docx'

# Карта: индекс параграфа → список тегов, которыми по порядку заменяются
# группы подчёркиваний (каждая группа «_____» = один слот для тега).
PARAGRAPH_REPLACEMENTS = {
    0:  ['{{ organization }}'],
    1:  ['{{ department }}'],
    4:  ['{{ permit_number }}'],
    6:  ['{{ work_manager_name }}', '{{ admitting_name }}'],
    7:  ['{{ producer_name }}',     '{{ observer_name }}'],
    8:  ['{{ brigade_line1 }}'],
    9:  ['{{ brigade_line2 }}'],
    10: ['{{ brigade_line3 }}'],
    11: ['{{ brigade_line4 }}'],
    12: ['{{ work_category }}'],
    13: ['{{ assignment_line1 }}'],
    14: ['{{ assignment_line2 }}'],
    15: ['{{ assignment_line3 }}'],
    16: ['{{ start_date }}',         '{{ start_time }}'],
    17: ['{{ end_date }}',           '{{ end_time }}'],
    18: ['{{ emergency_readiness_time }}'],
    22: ['{{ separate_instructions_line1 }}'],
    23: ['{{ separate_instructions_line2 }}'],
    24: ['{{ issuer_date }}',        '{{ issuer_time }}'],
    25: ['{{ issuer_signature }}',   '{{ issuer_last_name }}'],
    26: ['{{ extension_date }}',     '{{ extension_time }}'],
    27: ['{{ extension_signature }}','{{ extension_last_name }}'],
    31: ['{{ voltage_remains_at }}'],
    32: ['{{ admission_permit_signer }}', '{{ responsible_work_manager_signer }}'],
    43: ['{{ completion_notified_to }}', '{{ completion_date }}', '{{ completion_time }}'],
    45: ['{{ completion_producer_signature }}'],
    47: ['{{ completion_manager_signature }}'],
    49: ['{{ completion_admitter_signature }}'],
}


def replace_underscores(paragraph, tags):
    """Заменяет группы подчёркиваний в параграфе на теги по порядку."""
    full_text = paragraph.text

    # Спецслучай: параграфы 0/1 имеют шаблон «___ПРИМЕР___» — это одно поле
    # с демонстрационным примером, который надо удалить вместе с подчёркиваниями.
    # Эвристика: если между подчёркиваниями нет «/», значит это пример, а не label.
    # Сворачиваем все ___X___ , где X не содержит «/», в одну группу подчёркиваний.
    full_text = re.sub(r'_{2,}[^_/\n]+?_{2,}', '___', full_text, count=0)

    parts = re.split(r'_{2,}', full_text)
    if len(parts) - 1 < len(tags):
        print(f'  [!] Параграф ожидает {len(parts)-1} замен, передано {len(tags)} тегов')
    result = parts[0]
    for i, segment in enumerate(parts[1:], start=0):
        if i < len(tags):
            result += tags[i] + segment
        else:
            result += '_____' + segment

    # Очищаем все runs и записываем результат в первый run (стиль первого run сохраняется)
    if not paragraph.runs:
        paragraph.add_run(result)
        return
    paragraph.runs[0].text = result
    for run in paragraph.runs[1:]:
        run.text = ''


def fill_cell(cell, text):
    """Записывает текст в ячейку, очищая предыдущее содержимое."""
    # Очищаем все параграфы кроме первого
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    # Очищаем runs
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def main():
    os.makedirs(os.path.dirname(DST), exist_ok=True)

    print(f'Открываю исходник: {SRC}')
    doc = Document(SRC)

    # ===== ПАРАГРАФЫ =====
    print('\n=== РАЗМЕТКА ПАРАГРАФОВ ===')
    for idx, tags in PARAGRAPH_REPLACEMENTS.items():
        if idx >= len(doc.paragraphs):
            print(f'  [SKIP] Параграф {idx} вне диапазона')
            continue
        p = doc.paragraphs[idx]
        before = p.text.strip()[:60]
        replace_underscores(p, tags)
        after = p.text.strip()[:60]
        print(f'  [{idx:02d}] {before}... → {after}...')

    # ===== ТАБЛИЦА 0: Меры безопасности =====
    # 3-строчный паттерн docxtpl для row-cycle:
    #   R1 = только {%tr for ...%}  (строка превратится в {% for %} директиву)
    #   R2 = данные с {{ }}        (эта строка клонируется в цикле)
    #   R3 = только {%tr endfor %} (превратится в {% endfor %})
    print('\n=== ТАБЛИЦА 0: Меры безопасности → 3-row cycle ===')
    t0 = doc.tables[0]
    fill_cell(t0.rows[1].cells[0], '{%tr for m in safety_measures %}')
    fill_cell(t0.rows[1].cells[1], '')
    fill_cell(t0.rows[2].cells[0], '{{ m.installation_name }}')
    fill_cell(t0.rows[2].cells[1], '{{ m.action_required }}')
    fill_cell(t0.rows[3].cells[0], '{%tr endfor %}')
    fill_cell(t0.rows[3].cells[1], '')

    # ===== ТАБЛИЦА 1: Разрешение на допуск =====
    # Всего 3 строки (R0 header, R1, R2). Используем как 1 фиксированную запись
    # (одно разрешение на наряд) — поле admission. Цикл здесь излишен.
    print('=== ТАБЛИЦА 1: Разрешение на допуск → фиксированные поля ===')
    t1 = doc.tables[1]
    fill_cell(t1.rows[1].cells[1], '{{ admission_datetime }}')
    fill_cell(t1.rows[1].cells[2], '{{ admission_from_whom }}')
    fill_cell(t1.rows[1].cells[3], '{{ admission_signature }}')

    # ===== ТАБЛИЦА 2: Ежедневный допуск =====
    # R0-R2 — заголовки, R3=for, R4=данные, R5=endfor
    print('=== ТАБЛИЦА 2: Ежедневный допуск → 3-row cycle ===')
    t2 = doc.tables[2]
    fill_cell(t2.rows[3].cells[0], '{%tr for adm in daily_admissions %}')
    for ci in range(1, 6):
        fill_cell(t2.rows[3].cells[ci], '')
    fill_cell(t2.rows[4].cells[0], '{{ adm.workplace_name }}')
    fill_cell(t2.rows[4].cells[1], '{{ adm.admission_datetime }}')
    fill_cell(t2.rows[4].cells[2], '{{ adm.admitter_signature }}')
    fill_cell(t2.rows[4].cells[3], '{{ adm.producer_signature }}')
    fill_cell(t2.rows[4].cells[4], '{{ adm.end_datetime }}')
    fill_cell(t2.rows[4].cells[5], '{{ adm.producer_end_signature }}')
    fill_cell(t2.rows[5].cells[0], '{%tr endfor %}')
    for ci in range(1, 6):
        fill_cell(t2.rows[5].cells[ci], '')

    # ===== ТАБЛИЦА 3: Изменения в составе бригады =====
    # R0 — header, R1=for, R2=данные, R3=endfor
    print('=== ТАБЛИЦА 3: Изменения в бригаде → 3-row cycle ===')
    t3 = doc.tables[3]
    fill_cell(t3.rows[1].cells[0], '{%tr for ch in brigade_changes %}')
    for ci in range(1, 4):
        fill_cell(t3.rows[1].cells[ci], '')
    fill_cell(t3.rows[2].cells[0], '{{ ch.introduced_member }}')
    fill_cell(t3.rows[2].cells[1], '{{ ch.removed_member }}')
    fill_cell(t3.rows[2].cells[2], '{{ ch.datetime }}')
    fill_cell(t3.rows[2].cells[3], '{{ ch.authorized_by }}')
    fill_cell(t3.rows[3].cells[0], '{%tr endfor %}')
    for ci in range(1, 4):
        fill_cell(t3.rows[3].cells[ci], '')

    # ===== ТАБЛИЦА 4: Целевой инструктаж =====
    # R0 — header «Инструктаж провёл / Инструктаж получил»
    # R1: Issuer | sig | Resp.Manager | sig
    # R2: Admitter | sig | Members | sigs
    # R3: Resp.Manager (or Producer) | sig | Producer + Members | sigs
    print('=== ТАБЛИЦА 4: Целевой инструктаж → фиксированные поля ===')
    t4 = doc.tables[4]
    # R1
    fill_cell(t4.rows[1].cells[1], '{{ briefing_conducted_by_issuer }}')
    fill_cell(t4.rows[1].cells[3], '{{ briefing_received_by_manager }}')
    # R2
    fill_cell(t4.rows[2].cells[1], '{{ briefing_conducted_by_admitter }}')
    fill_cell(t4.rows[2].cells[3], '{{ briefing_received_by_members }}')
    # R3
    fill_cell(t4.rows[3].cells[1], '{{ briefing_conducted_by_manager }}')
    fill_cell(t4.rows[3].cells[3], '{{ briefing_received_by_producer }}')

    # ===== СОХРАНЕНИЕ =====
    doc.save(DST)
    print(f'\n✅ Шаблон сохранён: {DST}')
    print(f'   Размер: {os.path.getsize(DST)} байт')


if __name__ == '__main__':
    main()
